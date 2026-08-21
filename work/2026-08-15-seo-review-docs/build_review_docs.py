import json
import re
from html import unescape
from html.parser import HTMLParser
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE = Path(__file__).resolve().parent / 'source'
OUTPUT = Path(__file__).resolve().parent / 'output'
BLUE = '2E74B5'
NAVY = '1F4D78'
HEADER_FILL = 'E8EEF5'
TABLE_WIDTH_DXA = 9360


def set_run_font(run, name='Calibri', size=11, color=None, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for margin, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{margin}'))
        if node is None:
            node = OxmlElement(f'w:{margin}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_row_rules(row, header=False):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement('w:cantSplit')
    tr_pr.append(cant_split)
    if header:
        tbl_header = OxmlElement('w:tblHeader')
        tbl_header.set(qn('w:val'), 'true')
        tr_pr.append(tbl_header)


def set_table_width(table, width_dxa=TABLE_WIDTH_DXA):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(width_dxa))
    tbl_w.set(qn('w:type'), 'dxa')
    layout = tbl_pr.find(qn('w:tblLayout'))
    if layout is None:
        layout = OxmlElement('w:tblLayout')
        tbl_pr.append(layout)
    layout.set(qn('w:type'), 'fixed')
    indent = tbl_pr.find(qn('w:tblInd'))
    if indent is None:
        indent = OxmlElement('w:tblInd')
        tbl_pr.append(indent)
    indent.set(qn('w:w'), '120')
    indent.set(qn('w:type'), 'dxa')


def set_cell_width(cell, width_dxa):
    cell.width = Inches(width_dxa / 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(width_dxa))
    tc_w.set(qn('w:type'), 'dxa')


class BodyParser(HTMLParser):
    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.events = []
        self.block_tag = None
        self.block_parts = []
        self.list_tag = None
        self.table_rows = None
        self.current_row = None
        self.current_cell = None
        self.current_cell_tag = None
        self.anchor_url = None

    def _text_target(self):
        if self.current_cell is not None:
            return self.current_cell
        return self.block_parts

    def _flush_block(self):
        if self.block_tag and self.block_parts:
            text = ''.join(self.block_parts).strip()
            if text:
                if self.block_tag in ('h2', 'h3'):
                    self.events.append(('heading', 2 if self.block_tag == 'h2' else 3, text))
                elif self.block_tag == 'li':
                    self.events.append(('list', self.list_tag or 'ul', text))
                else:
                    self.events.append(('paragraph', text))
        self.block_tag = None
        self.block_parts = []

    def handle_starttag(self, tag, attrs):
        attrs = dict(attrs)
        if tag in ('p', 'h2', 'h3', 'li'):
            self._flush_block()
            self.block_tag = tag
        elif tag in ('ul', 'ol'):
            self._flush_block()
            self.list_tag = tag
        elif tag == 'table':
            self._flush_block()
            self.table_rows = []
        elif tag == 'tr':
            self.current_row = []
        elif tag in ('th', 'td'):
            self.current_cell = []
            self.current_cell_tag = tag
        elif tag == 'a':
            self.anchor_url = attrs.get('href')
        elif tag == 'br':
            self._text_target().append('\n')

    def handle_endtag(self, tag):
        if tag == 'a':
            if self.anchor_url:
                self._text_target().append(f' [{self.anchor_url}]')
            self.anchor_url = None
        elif tag in ('th', 'td'):
            if self.current_row is not None and self.current_cell is not None:
                self.current_row.append(''.join(self.current_cell).strip())
            self.current_cell = None
            self.current_cell_tag = None
        elif tag == 'tr':
            if self.table_rows is not None and self.current_row:
                self.table_rows.append(self.current_row)
            self.current_row = None
        elif tag == 'table':
            self._flush_block()
            if self.table_rows:
                self.events.append(('table', self.table_rows))
            self.table_rows = None
        elif tag == 'li':
            self._flush_block()
        elif tag in ('p', 'h2', 'h3'):
            self._flush_block()
        elif tag in ('ul', 'ol'):
            self._flush_block()
            self.list_tag = None

    def handle_data(self, data):
        self._text_target().append(data)


def add_paragraph(doc, text='', style=None, bold=False, italic=False, color=None, size=11):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return paragraph


def add_hyperlink(paragraph, text, url):
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), paragraph.part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True))
    run = OxmlElement('w:r')
    r_pr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), BLUE)
    r_pr.append(color)
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement('w:t')
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f'Heading {level}')
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    if level == 1:
        set_run_font(run, size=16, color=BLUE, bold=True)
    elif level == 2:
        set_run_font(run, size=13, color=BLUE, bold=True)
    else:
        set_run_font(run, size=12, color=NAVY, bold=True)
    return paragraph


def add_bullet(doc, text, numbered=False):
    style = 'List Number' if numbered else 'List Bullet'
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_simple_table(doc, rows, header=True, widths=None):
    if not rows:
        return
    columns = max(len(row) for row in rows)
    table = doc.add_table(rows=0, cols=columns)
    table.style = 'Table Grid'
    set_table_width(table)
    if widths is None:
        widths = [TABLE_WIDTH_DXA // columns] * columns
    for row_index, row_data in enumerate(rows):
        row = table.add_row()
        set_row_rules(row, header=header and row_index == 0)
        for col_index in range(columns):
            cell = row.cells[col_index]
            set_cell_width(cell, widths[col_index] if col_index < len(widths) else TABLE_WIDTH_DXA // columns)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            value = row_data[col_index] if col_index < len(row_data) else ''
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.1
            run = paragraph.add_run(str(value))
            set_run_font(run, size=9.5, bold=header and row_index == 0)
            if header and row_index == 0:
                set_cell_shading(cell, HEADER_FILL)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_html_body(doc, html):
    parser = BodyParser()
    parser.feed(html)
    previous_list_type = None
    for event in parser.events:
        kind = event[0]
        if kind == 'heading':
            add_heading(doc, event[2], event[1])
        elif kind == 'paragraph':
            text = unescape(event[1])
            paragraph = add_paragraph(doc, text)
            if '\n' in text:
                paragraph.paragraph_format.space_after = Pt(6)
        elif kind == 'list':
            add_bullet(doc, unescape(event[2]), numbered=event[1] == 'ol')
            previous_list_type = event[1]
        elif kind == 'table':
            add_simple_table(doc, [[unescape(cell) for cell in row] for row in event[1]])
            previous_list_type = None


def add_label_value(doc, label, value):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.2
    label_run = paragraph.add_run(f'{label}：')
    set_run_font(label_run, bold=True, color=NAVY)
    value_run = paragraph.add_run(str(value))
    set_run_font(value_run)
    return paragraph


def add_url_list(doc, links):
    for label, url in links:
        paragraph = doc.add_paragraph(style='List Bullet')
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(f'{label}：')
        set_run_font(run, bold=True)
        add_hyperlink(paragraph, url, url)


def add_document_header_footer(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run('Kat Chang 凱特營養師｜SEO 審閱稿')
    set_run_font(run, size=9, color='64748B')
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('僅供審閱，尚未代表文章通過審閱')
    set_run_font(run, size=9, color='64748B')


def configure_styles(doc):
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for level, size, color, before, after in [(1, 16, BLUE, 18, 10), (2, 13, BLUE, 14, 7), (3, 12, NAVY, 10, 5)]:
        style = doc.styles[f'Heading {level}']
        style.font.name = 'Calibri'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def build_doc(data, output_path):
    doc = Document()
    configure_styles(doc)
    add_document_header_footer(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run(data['reviewTitle'])
    set_run_font(run, size=20, color=NAVY, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Nutrition Concepts & Controversies 第 17 版書籍連載')
    set_run_font(run, size=11, color='64748B', italic=True)

    add_label_value(doc, '文件用途', '供人工審閱的完整 SEO 草稿與研究回報，不代表已通過審閱。')
    add_label_value(doc, '作者', data['author'])

    add_heading(doc, 'SEO 標題', 1)
    add_paragraph(doc, data['seoTitle'])
    add_heading(doc, '目標搜尋字詞、相關搜尋字詞與搜尋意圖', 1)
    add_label_value(doc, '目標搜尋字詞', '、'.join(data['targetTerms']))
    add_label_value(doc, '相關搜尋字詞', '、'.join(data['relatedTerms']))
    add_label_value(doc, '搜尋意圖', data['searchIntent'])
    add_heading(doc, '文章摘要與適合搜尋結果顯示的開場', 1)
    add_label_value(doc, '文章摘要', data['summary'])
    add_label_value(doc, '開場', data['opening'])

    add_heading(doc, '正文', 1)
    add_heading(doc, data['articleTitle'], 2)
    add_html_body(doc, data['bodyHtml'])

    add_heading(doc, 'SEO 描述', 1)
    add_paragraph(doc, data['seoDescription'])
    add_heading(doc, '文章分類、標籤與網址 slug', 1)
    add_label_value(doc, '分類', data['category'])
    add_label_value(doc, '標籤', '、'.join(data['tags']))
    add_label_value(doc, '網址 slug', data['slug'])
    add_label_value(doc, 'canonical', data['canonical'])

    add_heading(doc, '594katchang-source.github.io 站內連結建議', 1)
    add_url_list(doc, data['internalLinks'])

    add_heading(doc, 'FAQ 題目與結構化資料建議', 1)
    for question in data['faq']:
        add_bullet(doc, question)
    add_label_value(doc, 'FAQPage 建議', data['faqSchema'])
    add_heading(doc, '文章結構化資料、作者、更新日期與 canonical 建議', 1)
    add_label_value(doc, '文章結構化資料', '、'.join(data['articleSchema']))
    add_label_value(doc, '作者資料', data['author'])
    add_label_value(doc, '建議更新日期', data.get('reviewDate', '2026-08-15'))
    add_label_value(doc, 'canonical', data['canonical'])

    add_heading(doc, '來源連結與各來源支持的段落或主張', 1)
    source_rows = [['來源', '連結或檔案', '支持內容']]
    for label, url, scope in data['sources']:
        source_rows.append([label, url, scope])
    add_simple_table(doc, source_rows, widths=[2100, 3600, 3660])

    add_heading(doc, '文章字數、原創差異化主張與待確認事項', 1)
    add_label_value(doc, '正文可見字數', f"{data['wordCount']['characters']} 字，空白分隔詞數 {data['wordCount']['words']}。")
    add_heading(doc, '原創差異化主張', 2)
    for claim in data['originalClaims']:
        add_bullet(doc, claim)
    add_heading(doc, '待確認事項', 2)
    for item in data['pending']:
        add_bullet(doc, item)

    doc.core_properties.title = data['reviewTitle']
    doc.core_properties.subject = 'Kat Chang SEO 文章審閱檔'
    doc.core_properties.author = 'Kat Chang 凱特營養師'
    doc.core_properties.comments = '依 compact_reference_guide 版型建立，供人工審閱。'
    doc.save(output_path)


def main():
    OUTPUT.mkdir(parents=True, exist_ok=True)
    for name, filename in [
        ('chapter-1-review.json', 'chapter-01-food-choices-seo-review.docx'),
        ('chapter-2-review.json', 'chapter-02-nutrition-tools-seo-review.docx'),
    ]:
        data = json.loads((SOURCE / name).read_text(encoding='utf-8'))
        build_doc(data, OUTPUT / filename)
        print(f'{filename}: {OUTPUT / filename}')


if __name__ == '__main__':
    main()
