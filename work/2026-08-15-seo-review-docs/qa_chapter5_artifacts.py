import json
import re
from pathlib import Path
from zipfile import ZipFile
from xml.etree import ElementTree as ET

from docx import Document


BASE = Path(__file__).resolve().parent
SOURCE = BASE / "source" / "chapter-05-review.json"
HTML = BASE / "source" / "chapter-05-review.html"
DOCX = BASE / "output" / "chapter-05-lipids-seo-review.docx"
NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


def main():
    data = json.loads(SOURCE.read_text(encoding="utf-8"))
    html = HTML.read_text(encoding="utf-8")
    assert DOCX.exists() and DOCX.stat().st_size > 0
    assert data["wordCount"]["characters"] >= 2000
    assert data["bodyHtml"] in html
    assert f'<link rel="canonical" href="{data["canonical"]}">' in html
    assert len(re.findall(r"<h2>", data["bodyHtml"])) >= 10
    assert len(re.findall(r"<h3>", data["bodyHtml"])) == len(data["faqEntities"])
    assert len(re.findall(r"<table>", data["bodyHtml"])) >= 8
    for label, url in data["internalLinks"]:
        assert url in data["bodyHtml"] and url in html, url
    for item in data["faqEntities"]:
        assert item["question"] in data["bodyHtml"] and item["answer"] in data["bodyHtml"]

    with ZipFile(DOCX) as zf:
        assert "word/document.xml" in zf.namelist()
        root = ET.fromstring(zf.read("word/document.xml"))
        rel_root = ET.fromstring(zf.read("word/_rels/document.xml.rels"))
    rel_map = {item.get("Id"): item.get("Target") for item in rel_root}
    hyperlink_targets = []
    for hyperlink in root.findall(".//w:hyperlink", NS):
        rid = hyperlink.get(f"{{{REL_NS}}}id")
        if rid and rel_map.get(rid):
            hyperlink_targets.append(rel_map[rid])
    assert len(hyperlink_targets) == len(data["internalLinks"])
    for _, url in data["internalLinks"]:
        assert url in hyperlink_targets

    document = Document(DOCX)
    paragraph_text = "\n".join(p.text for p in document.paragraphs)
    required = [
        data["reviewTitle"], data["seoTitle"], "正文", "FAQ 題目與結構化資料建議",
        "來源連結與各來源支持的段落或主張", "原創差異化主張", "待確認事項",
        "Kat Chang 營養師的判讀", "七日脂質練習表：把「少油」改成可觀察的行動",
    ]
    for item in required:
        assert item in paragraph_text, item
    for forbidden in ["先說答案", "先給答案", "Chapter 5"]:
        assert forbidden not in paragraph_text, forbidden
    assert len(document.tables) >= 9

    table_audit = []
    for table in root.findall(".//w:tbl", NS):
        tbl_pr = table.find("./w:tblPr", NS)
        width = tbl_pr.find("./w:tblW", NS)
        layout = tbl_pr.find("./w:tblLayout", NS)
        assert width is not None and width.get(f"{{{NS['w']}}}w") == "9360"
        assert layout is not None and layout.get(f"{{{NS['w']}}}type") == "fixed"
        rows = table.findall("./w:tr", NS)
        assert rows and rows[0].find("./w:trPr/w:tblHeader", NS) is not None
        for row in rows:
            assert row.find("./w:trPr/w:cantSplit", NS) is not None
        table_audit.append(len(rows))

    print(json.dumps({
        "source": str(SOURCE),
        "html": str(HTML),
        "docx": str(DOCX),
        "bodyCharacters": data["wordCount"]["characters"],
        "bodyWords": data["wordCount"]["words"],
        "bodyTables": len(re.findall(r"<table>", data["bodyHtml"])),
        "bodyH2": len(re.findall(r"<h2>", data["bodyHtml"])),
        "bodyH3": len(re.findall(r"<h3>", data["bodyHtml"])),
        "faq": len(data["faqEntities"]),
        "internalLinks": len(data["internalLinks"]),
        "docxParagraphs": len(document.paragraphs),
        "docxTables": len(document.tables),
        "docxTableRows": table_audit,
        "sameSourceBody": True,
        "zipValid": True,
        "presetGeometry": "9360 DXA, fixed layout, header rows, cantSplit",
        "visualRender": "未完成，当前运行环境没有 LibreOffice soffice.exe",
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
