import html
import json
import re
from pathlib import Path

import docx
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[3]
REVIEW_DOCX = ROOT / "work/2026-08-15-seo-review-docs/output/chapter-06-proteins-amino-acids-seo-review.docx"
OUT_DIR = Path(__file__).resolve().parents[1]
SOURCE_HTML = OUT_DIR / "source/chapter-06-publish.html"
PUBLISH_JSON = OUT_DIR / "source/chapter-06-publish.json"
MANIFEST = OUT_DIR / "render/chapter6-publish-manifest.json"

URL_RE = re.compile(r"https?://[^\s<>\]\)\"']+")


def linkify(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    placeholders = []

    def bracketed(match: re.Match[str]) -> str:
        token = f"\x00LINK{len(placeholders)}\x00"
        url = match.group(1)
        placeholders.append(f'<a href="{html.escape(url, quote=True)}">{html.escape(url)}</a>')
        return token

    text = re.sub(r"\[(https?://[^\]\s]+)\]", bracketed, text)
    escaped = html.escape(text, quote=False)

    def plain_url(match: re.Match[str]) -> str:
        url = match.group(0)
        return f'<a href="{html.escape(url, quote=True)}">{html.escape(url)}</a>'

    escaped = URL_RE.sub(plain_url, escaped)
    for i, value in enumerate(placeholders):
        escaped = escaped.replace(f"\x00LINK{i}\x00", value)
    return escaped.replace("\n", "<br>")


def paragraph_html(text: str) -> str:
    body = linkify(text)
    if body.startswith("省時版本："):
        body = "<strong>省時版本：</strong>" + body[len("省時版本：") :]
    return f"<p>{body}</p>"


def table_html(table: Table) -> str:
    rows = []
    for row_index, row in enumerate(table.rows):
        cells = []
        for cell in row.cells:
            cell_text = cell.text.replace("\r\n", "\n").replace("\r", "\n")
            cells.append(f"<{('th' if row_index == 0 else 'td')}>{linkify(cell_text)}</{('th' if row_index == 0 else 'td')}>")
        rows.append(f"<tr>{''.join(cells)}</tr>")
    if not rows:
        return ""
    return f"<table><thead>{rows[0]}</thead><tbody>{''.join(rows[1:])}</tbody></table>"


def read_metadata(paragraphs):
    values = {}
    for index, para in enumerate(paragraphs):
        text = para.text.strip()
        if text == "SEO 標題" and index + 1 < len(paragraphs):
            values["title"] = paragraphs[index + 1].text.strip()
        elif text.startswith("文章摘要："):
            values["excerpt"] = text.split("：", 1)[1].strip()
        elif text.startswith("搜尋意圖："):
            values["search_intent"] = text.split("：", 1)[1].strip()
        elif text.startswith("目標搜尋字詞："):
            values["keywords"] = [x.strip() for x in text.split("：", 1)[1].split("、") if x.strip()]
        elif text.startswith("網址 slug："):
            values["slug"] = text.split("：", 1)[1].strip()
        elif text.startswith("canonical："):
            values["canonical"] = text.split("：", 1)[1].strip()
        elif text.startswith("SEO 描述"):
            pass
        elif text.startswith("建議更新日期："):
            values["date"] = text.split("：", 1)[1].strip()
        elif text.startswith("作者："):
            values["author"] = text.split("：", 1)[1].strip()
        elif text.startswith("分類："):
            values["category"] = text.split("：", 1)[1].strip()
        elif text.startswith("標籤："):
            values["tags"] = [x.strip() for x in text.split("：", 1)[1].split("、") if x.strip()]
        elif text.startswith("本篇章節："):
            values["chapter"] = text.split("：", 1)[1].strip()
        elif text.startswith("本篇整理書籍："):
            values["book"] = text.split("：", 1)[1].strip()
    return values


def extract_body(doc):
    chunks = []
    paragraph_records = []
    active = False
    stopped = False
    open_list = None

    def flush_list():
        nonlocal open_list
        if open_list:
            chunks.append(f"<{open_list[0]}>{''.join(open_list[1])}</{open_list[0]}>")
            open_list = None

    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            para = Paragraph(child, doc)
            text = para.text.strip()
            style = para.style.name if para.style else ""
            paragraph_records.append((text, style))
            if style == "Heading 1" and text == "正文":
                active = True
                continue
            if style == "Heading 1" and text == "SEO 描述":
                flush_list()
                stopped = True
                break
            if not active or not text:
                continue
            if style == "Heading 2":
                flush_list()
                chunks.append(f"<h2>{html.escape(text)}</h2>")
            elif style == "Heading 3":
                flush_list()
                chunks.append(f"<h3>{html.escape(text)}</h3>")
            elif style == "List Number":
                if open_list is None or open_list[0] != "ol":
                    flush_list()
                    open_list = ["ol", []]
                open_list[1].append(f"<li>{linkify(text)}</li>")
            elif style == "List Bullet":
                if open_list is None or open_list[0] != "ul":
                    flush_list()
                    open_list = ["ul", []]
                open_list[1].append(f"<li>{linkify(text)}</li>")
            else:
                flush_list()
                chunks.append(paragraph_html(text))
        elif child.tag == qn("w:tbl") and active and not stopped:
            flush_list()
            chunks.append(table_html(Table(child, doc)))
    flush_list()
    return "\n".join(chunks), paragraph_records


def extract_faq(doc):
    paragraphs = doc.paragraphs
    in_faq = False
    items = []
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        style = para.style.name if para.style else ""
        if style == "Heading 2" and text.startswith("FAQ："):
            in_faq = True
            continue
        if in_faq and style == "Heading 2" and not text.startswith("FAQ："):
            break
        if in_faq and style == "Heading 3" and i + 1 < len(paragraphs):
            answer = paragraphs[i + 1].text.strip()
            if answer and paragraphs[i + 1].style.name not in {"Heading 2", "Heading 3"}:
                items.append({"question": text, "answer": answer})
    return items


def source_rows(doc):
    if len(doc.tables) < 12:
        return []
    return [[cell.text.strip() for cell in row.cells] for row in doc.tables[11].rows]


def main():
    OUT_DIR.joinpath("source").mkdir(parents=True, exist_ok=True)
    doc = docx.Document(REVIEW_DOCX)
    metadata = read_metadata(doc.paragraphs)
    body, records = extract_body(doc)
    faq = extract_faq(doc)
    post_id = metadata["slug"]
    post = {
        "id": post_id,
        "title": metadata["title"],
        "date": metadata["date"],
        "excerpt": metadata["excerpt"],
        "keywords": metadata["keywords"],
        "showOnHome": False,
        "body": body,
        "faq": faq,
        "category": metadata.get("category"),
        "tags": metadata.get("tags", []),
        "canonical": metadata.get("canonical"),
        "author": metadata.get("author"),
        "dateModified": metadata.get("date"),
    }
    PUBLISH_JSON.write_text(json.dumps(post, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    SOURCE_HTML.write_text(body + "\n", encoding="utf-8")
    visible = re.sub(r"<[^>]+>", "", body)
    visible = re.sub(r"\s+", "", visible)
    manifest = {
        "source_docx": str(REVIEW_DOCX),
        "source_docx_sha256": "1133CEA9F1BDACAD8F6077BF1C28ED3A6ED65ED9A563C7093E685B6F3C09F3A2",
        "post_id": post_id,
        "title": post["title"],
        "date": post["date"],
        "visible_characters": len(visible),
        "body_html_bytes": len(body.encode("utf-8")),
        "faq_count": len(faq),
        "body_table_count": body.count("<table>"),
        "source_rows": source_rows(doc),
        "paragraph_record_count": len(records),
    }
    MANIFEST.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
