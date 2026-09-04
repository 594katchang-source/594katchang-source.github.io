import html
import re
import sys
from html.parser import HTMLParser
from pathlib import Path

from docx import Document

sys.stdout.reconfigure(encoding='utf-8')

ROOT = Path(__file__).resolve().parents[3]
POSTS = ROOT / 'blog' / 'posts.json'
OUTPUT = ROOT / 'work' / '2026-08-15-seo-review-docs' / 'output'
MAPPING = {
    '2026-08-14-food-choices-human-health-guide': 'chapter-01-food-choices-seo-review.docx',
    '2026-08-15-nutrition-tools-standards-guidelines': 'chapter-02-nutrition-tools-seo-review.docx',
    '2026-08-16-remarkable-body-nutrition-guide': 'chapter-03-remarkable-body-seo-review.docx',
    '2026-08-17-carbohydrates-food-guide': 'chapter-04-carbohydrates-seo-review.docx',
    '2026-08-20-lipids-fatty-acids-guide': 'chapter-05-lipids-seo-review.docx',
    '2026-08-22-proteins-amino-acids-book-notes': 'chapter-06-proteins-amino-acids-seo-review.docx',
    '2026-08-25-vitamins-book-notes': 'chapter-07-vitamins-seo-review.docx',
}


class PlainHTML(HTMLParser):
    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.parts = []

    def handle_data(self, data):
        self.parts.append(data)


def clean(value):
    return re.sub(r'\s+', ' ', html.unescape(value)).strip()


def html_text(value):
    parser = PlainHTML()
    parser.feed(value)
    return clean(' '.join(parser.parts))


def docx_text(path):
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return clean(' '.join(parts))


def sentences(value):
    return [clean(item) for item in re.split(r'(?<=[。！？])', value) if len(clean(item)) >= 18]


posts = __import__('json').loads(POSTS.read_text(encoding='utf-8'))['posts']
print('published_blog_to_word_comparison')
for post_id, docx_name in MAPPING.items():
    post = next(item for item in posts if item['id'] == post_id)
    word_path = OUTPUT / docx_name
    blog_text = html_text(post.get('body', ''))
    word_text = docx_text(word_path)
    units = sentences(blog_text)
    matched = [unit for unit in units if unit in word_text]
    missing = [unit for unit in units if unit not in word_text]
    print(f'ID: {post_id}')
    print(f'  blog_body_chars: {len(blog_text)}')
    print(f'  word_all_text_chars: {len(word_text)}')
    print(f'  sentence_units: {len(units)} matched_in_word: {len(matched)} missing_from_word: {len(missing)}')
    for unit in missing[:4]:
        print(f'  missing: {unit[:180]}')
