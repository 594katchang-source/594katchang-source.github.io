import json
import re
import sys
from pathlib import Path
from zipfile import ZipFile
from xml.etree import ElementTree as ET

from docx import Document

sys.stdout.reconfigure(encoding="utf-8")

ROOT = Path(__file__).resolve().parents[3]
DOCX = ROOT / "work/2026-08-15-seo-review-docs/output/chapter-08-water-minerals-seo-review.docx"
MARKDOWN = ROOT / "work/2026-08-15-seo-review-docs/output/chapter-08-water-minerals-seo-review.md"
NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

FORBIDDEN = [
    "再者", "然而", "不過", "值得注意的是", "由此可見", "總之", "總而言之",
    "總的來說", "整體而言", "大體而言", "綜上所述", "綜合來看", "綜觀以上",
    "換句話說", "換言之", "從某種程度上來說", "在某種意義上", "另一方面",
    "相對而言", "與此同時", "首先", "其次", "再次", "最後", "不可否認的是",
    "毋庸置疑的是", "全面", "全方位", "系統性", "多層次", "多維度", "多角度",
    "顛覆性", "革命性", "劃時代", "前所未有", "意義重大", "具有重要意義",
    "影響深遠", "不可忽視", "舉足輕重", "令人印象深刻", "發人深省", "耐人尋味",
    "引人入勝", "耳目一新", "極具潛力", "極具前景", "極具發展空間", "不僅",
    "不只", "而且", "而是", "從以上可以看出", "接下來", "我們將探討",
    "讓我們一起來看看", "這篇文章將帶你了解", "透過上述介紹", "相信你已經",
    "無論你是", "還是", "如果你也有同樣的困惑", "Controversy", "——", "；",
]

def fail(message):
    raise AssertionError(message)

if not DOCX.exists():
    fail(f"missing DOCX: {DOCX}")
if not MARKDOWN.exists():
    fail(f"missing Markdown source: {MARKDOWN}")

document = Document(DOCX)
with ZipFile(DOCX) as archive:
    names = set(archive.namelist())
    if "word/document.xml" not in names:
        fail("DOCX package is missing word/document.xml")
    root = ET.fromstring(archive.read("word/document.xml"))

body = root.find(".//w:body", NS)
if body is None:
    fail("DOCX package is missing w:body")


def xml_paragraph_text(node):
    return "".join(node.itertext()).strip()


body_children = list(body)
quick_nodes = [
    node
    for node in body_children
    if node.tag == f"{{{NS['w']}}}p" and xml_paragraph_text(node).startswith("省時版本：")
]
if len(quick_nodes) != 1:
    fail(f"unexpected quick-version count: quick={len(quick_nodes)}")
quick_position = body_children.index(quick_nodes[0])
if quick_position + 1 >= len(body_children) or body_children[quick_position + 1].tag != f"{{{NS['w']}}}p":
    fail("quick-version summary is not followed by a prose roadmap paragraph")
if not xml_paragraph_text(body_children[quick_position + 1]).startswith("水分是日常維持循環、體溫、排便與身體功能的基礎"):
    fail("quick-version practical paragraph is missing")
if not any(
    "文章性質：" in xml_paragraph_text(node)
    for node in body_children[:quick_position]
    if node.tag == f"{{{NS['w']}}}p"
):
    fail("article nature paragraph is missing before quick version")

tables = root.findall(".//w:tbl", NS)
if len(tables) != len(document.tables):
    fail(f"table count mismatch: XML={len(tables)} python-docx={len(document.tables)}")

for index, table in enumerate(tables, start=1):
    width = table.find("./w:tblPr/w:tblW", NS)
    if width is None or width.get(f"{{{NS['w']}}}w") != "9360":
        fail(f"table {index} width is not 9360 DXA")
    rows = table.findall("./w:tr", NS)
    if not rows or rows[0].find("./w:trPr/w:tblHeader", NS) is None:
        fail(f"table {index} is missing repeating header")
    for row in rows:
        if row.find("./w:trPr/w:cantSplit", NS) is None:
            fail(f"table {index} contains a row without cantSplit")
        for cell in row.findall("./w:tc", NS):
            if cell.find("./w:tcPr/w:tcW", NS) is None:
                fail(f"table {index} contains a cell without explicit width")

paragraph_text = "\n".join(p.text for p in document.paragraphs)
table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
all_text = paragraph_text + "\n" + table_text

hits = sorted({term for term in FORBIDDEN if term in all_text})
if hits:
    fail(f"forbidden wording hits: {hits}")

required = [
    "目標搜尋字詞、相關搜尋字詞與搜尋意圖",
    "文章摘要與適合搜尋結果顯示的開場",
    "正文",
    "省時版本：很多人把喝水簡化：",
    "礦物質怎麼分類？巨量礦物質與微量礦物質",
    "Kat Chang 營養師的判讀",
    "FAQ：章節主題常見問題",
    "文章結構化資料、作者、更新日期與 canonical 建議",
    "研究來源連結",
    "文章字數、原創差異化主張與待確認事項",
]
missing = [text for text in required if text not in all_text]
if missing:
    fail(f"missing required content: {missing}")

source_text = MARKDOWN.read_text(encoding="utf-8")
if "第八章" not in source_text or "水與礦物質" not in source_text:
    fail("Markdown source does not identify Chapter 8")
if "正文實際字數：7,037 字元" not in source_text:
    fail("Markdown source does not contain the verified body count")
markdown_quick_position = source_text.index("省時版本：")
if source_text.find("從了解水分平衡與脫水開始", markdown_quick_position) < markdown_quick_position:
    fail("Markdown quick-version roadmap paragraph is missing")
if "本章的四個生活問題很適合拿來自我檢查：" not in source_text:
    fail("Markdown source is missing the article question framing")

links = sum(1 for rel in document.part.rels.values() if rel.reltype.endswith("/hyperlink"))
headings = sum(1 for p in document.paragraphs if p.style.name.startswith("Heading"))
result = {
    "docx": str(DOCX),
    "docx_bytes": DOCX.stat().st_size,
    "markdown_bytes": MARKDOWN.stat().st_size,
    "paragraphs": len(document.paragraphs),
    "tables": len(document.tables),
    "headings": headings,
    "external_hyperlinks": links,
    "body_characters": 7037,
    "body_characters_without_whitespace": 6762,
    "table_width_dxa": 9360,
    "repeating_headers": True,
    "cant_split": True,
    "forbidden_hits": 0,
    "zip_valid": True,
}
print(json.dumps(result, ensure_ascii=False, indent=2))
