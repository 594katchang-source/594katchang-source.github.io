import html
import json
import re
import sys
from pathlib import Path
from zipfile import ZipFile
from xml.etree import ElementTree as ET

from docx import Document

sys.stdout.reconfigure(encoding='utf-8')

ROOT = Path(__file__).resolve().parents[3]
POSTS = ROOT / 'blog' / 'posts.json'
OUTPUT = ROOT / 'work' / '2026-08-15-seo-review-docs' / 'output'
NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
MAPPING = {
    '2026-08-14-food-choices-human-health-guide': 'chapter-01-food-choices-seo-review.docx',
    '2026-08-15-nutrition-tools-standards-guidelines': 'chapter-02-nutrition-tools-seo-review.docx',
    '2026-08-16-remarkable-body-nutrition-guide': 'chapter-03-remarkable-body-seo-review.docx',
    '2026-08-17-carbohydrates-food-guide': 'chapter-04-carbohydrates-seo-review.docx',
    '2026-08-20-lipids-fatty-acids-guide': 'chapter-05-lipids-seo-review.docx',
    '2026-08-22-proteins-amino-acids-book-notes': 'chapter-06-proteins-amino-acids-seo-review.docx',
    '2026-08-25-vitamins-book-notes': 'chapter-07-vitamins-seo-review.docx',
}


def compact(value):
    return re.sub(r'\s+', '', value or '')


def html_first_h2(body):
    match = re.search(r'<h2>(.*?)</h2>', body, flags=re.S)
    return compact(html.unescape(match.group(1))) if match else ''


def paragraph_value(values, label):
    for index, value in enumerate(values):
        if value.strip() == label and index + 1 < len(values):
            return values[index + 1].strip()
    return ''


def audit(post, filename):
    path = OUTPUT / filename
    assert path.exists(), path
    document = Document(path)
    values = [paragraph.text for paragraph in document.paragraphs]
    title = paragraph_value(values, 'SEO 標題')
    excerpt = paragraph_value(values, '文章摘要與適合搜尋結果顯示的開場')
    keywords = paragraph_value(values, '目標搜尋字詞、相關搜尋字詞與搜尋意圖')
    category = next(value[3:].strip() for value in values if value.startswith('分類：'))
    date = next(value[7:].strip() for value in values if value.startswith('建議更新日期：'))
    assert title == post['title'], (filename, 'title')
    assert excerpt == '文章摘要：' + post['excerpt'], (filename, 'excerpt')
    assert keywords == '目標搜尋字詞：' + '、'.join(post.get('keywords', [])), (filename, 'keywords')
    assert category == post.get('category', ''), (filename, 'category')
    assert date == post.get('date', ''), (filename, 'date')
    all_text = '\n'.join(values)
    for table in document.tables:
        for row in table.rows:
            all_text += '\n' + '\n'.join(cell.text for cell in row.cells)
    assert compact(html_first_h2(post['body'])) in compact(all_text), (filename, 'body heading')
    with ZipFile(path) as archive:
        names = set(archive.namelist())
        assert 'word/document.xml' in names, (filename, 'document.xml')
        root = ET.fromstring(archive.read('word/document.xml'))
    tables = root.findall('.//w:tbl', NS)
    assert len(tables) == len(document.tables), (filename, 'table count')
    for table in tables:
        width = table.find('./w:tblPr/w:tblW', NS)
        assert width is not None and width.get('{%s}w' % NS['w']) == '9360', (filename, 'table width')
        rows = table.findall('./w:tr', NS)
        assert rows and rows[0].find('./w:trPr/w:tblHeader', NS) is not None, (filename, 'table header')
        for row in rows:
            assert row.find('./w:trPr/w:cantSplit', NS) is not None, (filename, 'cantSplit')
    return {
        'file': str(path),
        'sha256': __import__('hashlib').sha256(path.read_bytes()).hexdigest(),
        'paragraphs': len(document.paragraphs),
        'tables': len(document.tables),
        'metadata_match': True,
        'body_heading_match': True,
        'zip_valid': True,
        'table_geometry': '9360 DXA',
        'repeating_headers': True,
        'cant_split': True,
        'visual_render': '未完成，缺少 LibreOffice soffice.exe',
    }


posts = {post['id']: post for post in json.loads(POSTS.read_text(encoding='utf-8'))['posts']}
results = [audit(posts[post_id], filename) for post_id, filename in MAPPING.items()]
print(json.dumps(results, ensure_ascii=False, indent=2))
