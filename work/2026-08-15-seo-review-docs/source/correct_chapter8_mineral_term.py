from pathlib import Path

from docx import Document


DOCX = Path(__file__).resolve().parents[1] / "output" / "chapter-08-water-minerals-seo-review.docx"
OLD = "主要礦物質"
NEW = "巨量礦物質"


def main():
    document = Document(DOCX)
    matches = []
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    for paragraph in paragraphs:
        if OLD in paragraph.text:
            matches.append(paragraph)
            for run in paragraph.runs:
                run.text = run.text.replace(OLD, NEW)
    if len(matches) != 2:
        raise RuntimeError(f"Expected two remaining term matches, found {len(matches)}")
    heading = [p for p in document.paragraphs if p.text == "礦物質怎麼分類？巨量礦物質與微量礦物質"]
    if len(heading) != 1:
        raise RuntimeError(f"Expected corrected heading, found {len(heading)}")
    document.save(DOCX)
    print(f"Updated: {DOCX}")


if __name__ == "__main__":
    main()
