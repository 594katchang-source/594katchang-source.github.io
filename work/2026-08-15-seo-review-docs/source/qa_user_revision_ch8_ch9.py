from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

ROOT = Path(__file__).resolve().parents[3]
OUTPUT = ROOT / "work/2026-08-15-seo-review-docs/output"
SOURCE = ROOT / "work/2026-08-15-seo-review-docs/source"


def fail(message: str) -> None:
    raise AssertionError(message)


def blocks(document: Document):
    result = []
    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            result.append(Paragraph(child, document._body))
        elif child.tag.endswith("}tbl"):
            result.append(Table(child, document._body))
    return result


def block_text(block) -> str:
    if isinstance(block, Paragraph):
        return block.text
    return "\n".join(cell.text for row in block.rows for cell in row.cells)


def body_text(document: Document, start: str, end: str) -> list[str]:
    items = blocks(document)
    start_i = next(i for i, item in enumerate(items) if isinstance(item, Paragraph) and item.text.startswith(start))
    end_i = next(i for i, item in enumerate(items) if isinstance(item, Paragraph) and item.text == end)
    return [block_text(item) for item in items[start_i:end_i] if block_text(item).strip()]


def read_word(name: str) -> tuple[Document, str]:
    document = Document(OUTPUT / name)
    text = "\n".join(block_text(item) for item in blocks(document))
    return document, text


ch8_doc, ch8_word = read_word("chapter-08-water-minerals-seo-review.docx")
ch9_doc, ch9_word = read_word("chapter-09-energy-balance-seo-review.docx")
ch8_md = (OUTPUT / "chapter-08-water-minerals-seo-review.md").read_text(encoding="utf-8")
ch9_md = (OUTPUT / "chapter-09-energy-balance-seo-review.md").read_text(encoding="utf-8")
ch9_html = (SOURCE / "chapter-09-energy-balance-body.html").read_text(encoding="utf-8")

title_old = "每天喝多少水才夠？從電解質、鈣鐵到骨質保養的生活判讀"
title_new = "每天喝多少水才夠？從電解質、鈣鐵到骨質保養"
for label, text in [("Chapter 8 Word", ch8_word), ("Chapter 8 Markdown", ch8_md)]:
    if title_new not in text:
        fail(f"{label} is missing the concise title")
    if title_old in text:
        fail(f"{label} still contains the old title")

ch8_quick = [p for p in ch8_doc.paragraphs if p.text.startswith("省時版本：")]
if len(ch8_quick) != 1:
    fail(f"Chapter 8 quick-version paragraph count: {len(ch8_quick)}")
if not ch8_quick[0].text.startswith("省時版本：很多人把喝水簡化："):
    fail("Chapter 8 quick-version opening does not follow the reviewed format")
ch8_quick_index = next(i for i, p in enumerate(ch8_doc.paragraphs) if p.text.startswith("省時版本："))
if len(ch8_quick) != 1 or not any(
    p.text.startswith("水分是日常維持循環、體溫、排便與身體功能的基礎。")
    for p in ch8_doc.paragraphs[ch8_quick_index + 1 :]
):
    fail("Chapter 8 quick-version practical paragraph is missing")
if "從了解水分平衡與脫水開始" not in ch8_word:
    fail("Chapter 8 quick-version roadmap is missing from Word")
if "| 你想處理的問題 |" in ch8_md or "## 省時版本：" in ch8_md:
    fail("Chapter 8 still uses the old quick-version table format")
if "省時版本：很多人把喝水簡化：" not in ch8_md:
    fail("Chapter 8 Markdown quick-version prose is missing")

ch9_quick = [p for p in ch9_doc.paragraphs if p.text.startswith("省時版本：")]
if len(ch9_quick) != 1:
    fail(f"Chapter 9 quick-version paragraph count: {len(ch9_quick)}")
if "從能量平衡與體重變化開始" not in ch9_word:
    fail("Chapter 9 quick-version roadmap is missing from Word")
if "| 你現在想處理的問題 |" in ch9_md or "## 省時版本：" in ch9_md:
    fail("Chapter 9 still uses the old quick-version table format")
if "省時版本：很多人把減重理解成一條算式" not in ch9_md:
    fail("Chapter 9 Markdown quick-version prose is missing")
if "<h2>省時版本：</h2>" in ch9_html or "<table><tr><th>你現在想處理的問題" in ch9_html:
    fail("Chapter 9 same-source HTML still uses the old quick-version table")

ch8_new_phrases = [
    "礦物質怎麼分類？巨量礦物質與微量礦物質",
    "接著認識電解質與巨量礦物質",
    "這是一個日常起點建議，不適用於心臟衰竭",
    "口渴、尿液顏色與脫水，怎麼看？",
    "尿液顏色可作為生活觀察線索，但可能會受",
    "問題已經超出一般飲水技巧可以處理的範圍，需要迅速接受醫療評估",
    "這時不能只靠白開水自行補充",
    "這裡建議的原則是：",
    "電解質是什麼？水分分布為什麼會受溶質影響？",
    "不是只靠少吃食鹽、控制食鹽而已。",
    "鈣、鎂、磷的作用彼此相關，但不能把它們當成一組買來就能一次補足",
    "血液中的鈣正常，不代表骨骼的鈣儲備充足。",
    "但疲倦也可能來自睡眠不足",
    "健康者則無需避開茶或乳品",
    "怎麼安排每天都吃得到鈣來源？",
    "把鈣分布在餐食，比每天睡前才想起一顆鈣片更容易持續、益處更多。",
    "下表是依據生活作息建議，不代表固定份量。",
    "水和礦物質這部分可以分成三個重點。",
    "補充水分是日常基礎，分次飲用比追逐某個目標數字有用。",
    "綜合評估三天飲食、用藥、保健品、症狀與檢驗結果",
    "這些事情看似很普通，卻比單一神奇食物更值得你做。",
    "營養要跟治療計畫一起安排",
]
for phrase in ch8_new_phrases:
    if phrase not in ch8_word or phrase not in ch8_md:
        fail(f"Chapter 8 missing revised phrase: {phrase}")

ch8_old_phrases = [
    title_old,
    "口渴、尿液顏色與脫水，怎麼一起看？",
    "尿液顏色可作為生活觀察線索，卻會被",
    "這裡的實用原則很簡單：",
    "電解質是什麼？為什麼水會跟著鹽走？",
    "單靠少撒一點鹽，常還不夠。",
    "三者功能相連，卻不代表購買一瓶「鈣鎂磷」就能補好。",
    "血液中的鈣正常，不能直接證明骨骼儲備充足。",
    "健康者無需把茶或乳品全部刪除",
    "每天吃得到的鈣來源怎麼安排？",
    "把鈣塞進三餐，比每天睡前才想起一顆鈣片更容易持續。",
    "下表是生活安排，不代表每個人的固定份量。",
    "我會把這一章轉成三個生活順序。",
    "主要礦物質",
    "白開水是日常基礎，分次飲用比追逐一個漂亮數字實用。",
    "這些事情看似平常，卻比單一神奇食物更值得投入。",
    "營養只能放在治療計畫裡一起安排",
]
for phrase in ch8_old_phrases:
    if phrase in ch8_word or phrase in ch8_md:
        fail(f"Chapter 8 still contains a disliked phrase: {phrase}")

ch8_body = body_text(ch8_doc, "開場：", "SEO 描述")
ch9_body = body_text(ch9_doc, "文章性質：", "SEO 描述")
for label, values, expected in [
    ("Chapter 8", ch8_body, ["開場：很多人把喝水簡化成一個數字", "電解質是什麼？", "鈣是骨骼所需營養素"]),
    ("Chapter 9", ch9_body, ["文章性質：依原章節概念順序整理", "BMI、腰圍與身體組成", "需要，但重點放在營養密度"]),
]:
    if not values or not values[0].startswith(expected[0]):
        fail(f"{label} first body sample is not preserved")
    middle = values[len(values) // 2]
    if expected[1] not in "\n".join(values):
        fail(f"{label} middle body sample is missing")
    if not values[-1].startswith(expected[2]):
        fail(f"{label} last body sample is not preserved")
    print(f"{label} first/middle/last samples: PASS; body_blocks={len(values)}, middle={middle[:80]}")

if "正文實際字數：7,037 字元" not in ch8_word or "正文實際字數：7,037 字元" not in ch8_md:
    fail("Chapter 8 Word/Markdown count metadata is not synchronized")
if "正文可見字數：5369 字" not in ch9_word or "正文實際字數：5369 字元" not in ch9_md:
    fail("Chapter 9 Word/Markdown count metadata is not synchronized")

print("user revision checks: PASS")
