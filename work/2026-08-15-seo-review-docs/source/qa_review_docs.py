import json
import re
from pathlib import Path
from zipfile import ZipFile
from xml.etree import ElementTree as ET

from docx import Document


BASE = Path(__file__).resolve().parent.parent
SOURCE = BASE / 'source'
OUTPUT = BASE / 'output'
NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


def qa_one(json_name, docx_name, expected_links):
    data = json.loads((SOURCE / json_name).read_text(encoding='utf-8'))
    docx_path = OUTPUT / docx_name
    assert docx_path.exists(), docx_path
    with ZipFile(docx_path) as zf:
        names = set(zf.namelist())
        assert 'word/document.xml' in names
        xml = ET.fromstring(zf.read('word/document.xml'))
    document = Document(docx_path)
    paragraphs = '\n'.join(p.text for p in document.paragraphs)
    table_count = len(document.tables)
    table_xml = xml.findall('.//w:tbl', NS)
    assert table_count >= 5, table_count
    assert len(table_xml) == table_count, (len(table_xml), table_count)
    for table in table_xml:
        width = table.find('./w:tblPr/w:tblW', NS)
        assert width is not None and width.get('{%s}w' % NS['w']) == '9360', width.attrib if width is not None else None
        rows = table.findall('./w:tr', NS)
        assert rows, 'table has no rows'
        assert rows[0].find('./w:trPr/w:tblHeader', NS) is not None, 'missing repeated header'
        for row in rows:
            assert row.find('./w:trPr/w:cantSplit', NS) is not None, 'missing cantSplit'
    for required in [data['reviewTitle'], '正文', 'FAQ 題目與結構化資料建議', '來源連結與各來源支持的段落或主張', '原創差異化主張', '待確認事項']:
        assert required in paragraphs, required
    for url in expected_links:
        assert url in paragraphs, url
    for forbidden in ['先說答案', '先給答案', 'Controversy', 'Chapter 1', 'Chapter 2']:
        assert forbidden not in paragraphs, forbidden
    assert data['wordCount']['characters'] >= 2000, data['wordCount']
    return {
        'file': str(docx_path),
        'paragraphs': len(document.paragraphs),
        'tables': table_count,
        'bodyCharacters': data['wordCount']['characters'],
        'linksChecked': len(expected_links),
        'zipValid': True,
        'visualRender': '未完成，缺少 LibreOffice soffice.exe'
    }


def main():
    links1 = [
        'https://594katchang-source.github.io/blog/post.html?id=2026-08-13-nutrition-concepts-controversies-17e-guide',
        'https://594katchang-source.github.io/blog/post.html?id=sample-balanced-breakfast',
        'https://594katchang-source.github.io/about.html',
        'https://594katchang-source.github.io/blog/',
        'https://594katchang-source.github.io/teach/paper-radar/'
    ]
    links2 = [
        'https://594katchang-source.github.io/blog/post.html?id=2026-08-13-nutrition-concepts-controversies-17e-guide',
        'https://594katchang-source.github.io/blog/post.html?id=2026-08-14-food-choices-human-health-guide',
        'https://594katchang-source.github.io/blog/',
        'https://594katchang-source.github.io/about.html',
        'https://594katchang-source.github.io/teach/paper-radar/'
    ]
    result = [
        qa_one('chapter-1-review.json', 'chapter-01-food-choices-seo-review.docx', links1),
        qa_one('chapter-2-review.json', 'chapter-02-nutrition-tools-seo-review.docx', links2)
    ]
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == '__main__':
    main()
