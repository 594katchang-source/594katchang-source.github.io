import json
import re
import sys
from pathlib import Path
from zipfile import ZipFile
from xml.etree import ElementTree as ET

from docx import Document

sys.stdout.reconfigure(encoding='utf-8')

base = Path(__file__).resolve().parents[1]
source_json = base / 'source' / 'chapter-09-review.json'
docx_path = base / 'output' / 'chapter-09-energy-balance-seo-review.docx'
md_path = base / 'output' / 'chapter-09-energy-balance-seo-review.md'
ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
checks = {}

checks['docx_exists'] = docx_path.exists()
checks['markdown_exists'] = md_path.exists()
checks['json_exists'] = source_json.exists()
with ZipFile(docx_path) as zf:
    names = set(zf.namelist())
    checks['zip_has_document_xml'] = 'word/document.xml' in names
    checks['zip_has_relationships'] = 'word/_rels/document.xml.rels' in names
    checks['zip_has_content_types'] = '[Content_Types].xml' in names
    root = ET.fromstring(zf.read('word/document.xml'))

doc = Document(docx_path)
all_paragraphs = '\n'.join(p.text for p in doc.paragraphs)
tables = root.findall('.//w:tbl', ns)
quick_paragraphs = [p for p in doc.paragraphs if p.text.startswith('省時版本：')]
quick_index = next((i for i, p in enumerate(doc.paragraphs) if p.text.startswith('省時版本：')), None)
checks['quick_version_is_two_paragraphs'] = (
    len(quick_paragraphs) == 1 and
    quick_index is not None and
    quick_index + 1 < len(doc.paragraphs) and
    doc.paragraphs[quick_index + 1].text.startswith('能量平衡可以幫助理解體重變化') and
    '從能量平衡與體重變化開始' in doc.paragraphs[quick_index + 1].text and
    any(p.text.startswith('文章性質：') for p in doc.paragraphs[:quick_index])
)
checks['quick_version_has_no_table'] = not any(
    row.cells[0].text.startswith('你現在想處理的問題')
    for table in doc.tables for row in table.rows[:1]
)
checks['table_count_matches'] = len(tables) == len(doc.tables)
checks['at_least_five_tables'] = len(tables) >= 5
checks['all_tables_9360_dxa'] = all(
    table.find('./w:tblPr/w:tblW', ns) is not None and
    table.find('./w:tblPr/w:tblW', ns).get('{%s}w' % ns['w']) == '9360'
    for table in tables
)
checks['all_rows_cant_split'] = all(
    row.find('./w:trPr/w:cantSplit', ns) is not None
    for table in tables for row in table.findall('./w:tr', ns)
)
checks['all_table_headers_repeat'] = all(
    table.findall('./w:tr', ns) and
    table.findall('./w:tr', ns)[0].find('./w:trPr/w:tblHeader', ns) is not None
    for table in tables
)
checks['required_sections'] = all(item in all_paragraphs for item in [
    '第九章待審 SEO 草稿', '正文', 'FAQ 題目與結構化資料建議',
    '來源連結與各來源支持的段落或主張', '原創差異化主張', '待確認事項'
])
checks['no_prior_chapter_contamination'] = all(item not in all_paragraphs for item in ['Chapter 1', 'Chapter 2', 'Chapter 8'])
checks['has_hyperlinks'] = len(root.findall('.//w:hyperlink', ns)) >= 7
checks['has_headings'] = len(re.findall(r'第九章|SEO 標題|正文|FAQ', all_paragraphs)) >= 4

for key, value in checks.items():
    print(f'{key}: {"PASS" if value else "FAIL"}')
print(f'paragraphs: {len(doc.paragraphs)}')
print(f'tables: {len(tables)}')
print(f'hyperlinks: {len(root.findall(".//w:hyperlink", ns))}')
if not all(checks.values()):
    raise SystemExit(1)
