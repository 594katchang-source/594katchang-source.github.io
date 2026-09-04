import html
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


sys.stdout.reconfigure(encoding="utf-8")

ROOT = Path(__file__).resolve().parents[3]
WORK = ROOT / "work" / "2026-08-15-seo-review-docs"
DOCX = WORK / "output" / "chapter-08-water-minerals-seo-review.docx"
SOURCE_HTML = WORK / "source" / "chapter-08-publish.html"
SOURCE_JSON = WORK / "source" / "chapter-08-publish.json"
MANIFEST = WORK / "render" / "chapter-08-publish-manifest.json"

PUBLISH_DATE = "2026-09-01"
PUBLISH_CATEGORY = "書籍連載與營養知識"


def inline_html(paragraph):
    """Render Word text and external hyperlinks without losing user-edited link targets."""
    pieces = []
    children = list(paragraph._p)
    index = 0
    while index < len(children):
        child = children[index]
        if child.tag == qn("w:hyperlink"):
            rid = child.get(qn("r:id"))
            url = paragraph.part.rels[rid].target_ref if rid in paragraph.part.rels else ""
            text = "".join((node.text or "") for node in child.findall(".//" + qn("w:t")))
            if url:
                pieces.append(
                    f'<a href="{html.escape(url, quote=True)}" target="_blank" rel="noopener">'
                    f"{html.escape(text)}"
                    "</a>"
                )
            else:
                pieces.append(html.escape(text))
            index += 1
            continue
        if child.tag != qn("w:r"):
            index += 1
            continue

        is_field_begin = any(
            node.tag == qn("w:fldChar") and node.get(qn("w:fldCharType")) == "begin"
            for node in child
        )
        if is_field_begin:
            instruction = []
            display = []
            separated = False
            end_index = index
            for field_index in range(index, len(children)):
                field_run = children[field_index]
                if field_run.tag != qn("w:r"):
                    continue
                field_ended = False
                for node in field_run:
                    if node.tag == qn("w:instrText"):
                        instruction.append(node.text or "")
                    elif node.tag == qn("w:fldChar"):
                        field_type = node.get(qn("w:fldCharType"))
                        if field_type == "separate":
                            separated = True
                        elif field_type == "end":
                            field_ended = True
                    elif separated and node.tag == qn("w:t"):
                        display.append(node.text or "")
                    elif separated and node.tag == qn("w:tab"):
                        display.append(" ")
                    elif separated and node.tag == qn("w:br"):
                        display.append("<br>")
                if field_ended:
                    end_index = field_index
                    break
            instruction_text = "".join(instruction)
            display_text = "".join(display)
            match = re.search(r'HYPERLINK\s+"([^"]+)"', instruction_text, flags=re.IGNORECASE)
            if match and display_text:
                pieces.append(
                    f'<a href="{html.escape(match.group(1), quote=True)}" target="_blank" rel="noopener">'
                    f"{html.escape(display_text)}"
                    "</a>"
                )
            else:
                pieces.append(html.escape(display_text))
            index = end_index + 1
            continue
        for node in child:
            if node.tag == qn("w:t"):
                pieces.append(html.escape(node.text or ""))
            elif node.tag == qn("w:tab"):
                pieces.append(" ")
            elif node.tag == qn("w:br"):
                pieces.append("<br>")
        index += 1
    return "".join(pieces)


def paragraph_html(paragraph):
    body = inline_html(paragraph)
    if body.startswith("省時版本："):
        body = "<strong>省時版本：</strong>" + body[len("省時版本：") :]
    return f"<p>{body}</p>"


def cell_html(cell):
    parts = []
    for paragraph in cell.paragraphs:
        value = inline_html(paragraph)
        if value:
            parts.append(value)
    return "<br>".join(parts)


def table_html(table):
    rows = []
    for row_index, row in enumerate(table.rows):
        tag = "th" if row_index == 0 else "td"
        cells = "".join(f"<{tag}>{cell_html(cell)}</{tag}>" for cell in row.cells)
        rows.append(f"<tr>{cells}</tr>")
    if not rows:
        return ""
    return f"<table><thead>{rows[0]}</thead><tbody>{''.join(rows[1:])}</tbody></table>"


def metadata(paragraphs):
    result = {}
    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if text == "SEO 標題" and index + 1 < len(paragraphs):
            result["title"] = paragraphs[index + 1].text.strip()
        elif text == "SEO 描述" and index + 1 < len(paragraphs):
            result["seoDescription"] = paragraphs[index + 1].text.strip()
        for line in text.splitlines():
            line = line.strip()
            if line.startswith("文章摘要："):
                result["excerpt"] = line.split("：", 1)[1].strip()
            elif line.startswith("目標搜尋字詞："):
                result["keywords"] = [x.strip() for x in line.split("：", 1)[1].split("、") if x.strip()]
            elif line.startswith("標籤："):
                result["tags"] = [x.strip() for x in line.split("：", 1)[1].split("、") if x.strip()]
            elif line.startswith("網址 slug："):
                result["slug"] = line.split("：", 1)[1].strip()
    return result


def extract_body(document):
    chunks = []
    active = False
    stopped = False
    open_list = None

    def flush_list():
        nonlocal open_list
        if open_list:
            chunks.append(f"<{open_list[0]}>{''.join(open_list[1])}</{open_list[0]}>")
            open_list = None

    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            paragraph = Paragraph(child, document)
            text = paragraph.text.strip()
            style = paragraph.style.name if paragraph.style else ""
            if style == "Heading 1" and text == "正文":
                active = True
                continue
            if style == "Heading 1" and text == "SEO 描述":
                flush_list()
                stopped = True
                break
            if not active or not text:
                continue
            if style == "Heading 2":
                flush_list()
                chunks.append(f"<h2>{html.escape(text)}</h2>")
            elif style == "Heading 3":
                flush_list()
                chunks.append(f"<h3>{html.escape(text)}</h3>")
            elif style == "List Number":
                if open_list is None or open_list[0] != "ol":
                    flush_list()
                    open_list = ["ol", []]
                open_list[1].append(f"<li>{inline_html(paragraph)}</li>")
            elif style == "List Bullet":
                if open_list is None or open_list[0] != "ul":
                    flush_list()
                    open_list = ["ul", []]
                open_list[1].append(f"<li>{inline_html(paragraph)}</li>")
            else:
                flush_list()
                chunks.append(paragraph_html(paragraph))
        elif child.tag == qn("w:tbl") and active and not stopped:
            flush_list()
            chunks.append(table_html(Table(child, document)))
    flush_list()
    return "\n".join(chunks)


def extract_faq(document):
    items = []
    in_faq = False
    paragraphs = document.paragraphs
    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        style = paragraph.style.name if paragraph.style else ""
        if style == "Heading 2" and text.startswith("FAQ："):
            in_faq = True
            continue
        if in_faq and style == "Heading 1":
            break
        if in_faq and style == "Heading 3" and index + 1 < len(paragraphs):
            answer = paragraphs[index + 1].text.strip()
            if answer and not answer.startswith("SEO "):
                items.append({"question": text, "answer": answer})
    return items


def visible_text(body):
    return re.sub(r"\s+", "", html.unescape(re.sub(r"<[^>]+>", "", body)))


def main():
    document = Document(DOCX)
    values = metadata(document.paragraphs)
    body = extract_body(document)
    faq = extract_faq(document)
    if not values.get("title") or not values.get("slug") or not body:
        raise RuntimeError("Chapter 8 Word metadata or body is incomplete")
    post_id = f"{PUBLISH_DATE}-{values['slug']}"
    post = {
        "id": post_id,
        "title": values["title"],
        "date": PUBLISH_DATE,
        "category": PUBLISH_CATEGORY,
        "excerpt": values.get("excerpt", ""),
        "keywords": values.get("keywords", []),
        "showOnHome": False,
        "body": body,
        "faq": faq,
    }
    SOURCE_HTML.write_text(body + "\n", encoding="utf-8")
    SOURCE_JSON.write_text(json.dumps(post, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    manifest = {
        "source_docx": str(DOCX),
        "source_docx_sha256": __import__("hashlib").sha256(DOCX.read_bytes()).hexdigest(),
        "post_id": post_id,
        "title": post["title"],
        "date": post["date"],
        "category": post["category"],
        "visible_characters": len(visible_text(body)),
        "body_html_bytes": len(body.encode("utf-8")),
        "faq_count": len(faq),
        "body_table_count": body.count("<table>"),
        "body_external_link_count": body.count("<a href="),
    }
    MANIFEST.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
