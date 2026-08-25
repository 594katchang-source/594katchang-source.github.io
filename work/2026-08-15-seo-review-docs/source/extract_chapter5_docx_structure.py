import sys
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding="utf-8")

ROOT = Path(__file__).resolve().parent.parent
DOCX = ROOT / "output" / "chapter-05-lipids-seo-review.docx"

def iter_block_items(parent):
    parent_elm = parent.element.body
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield "P", next((p for p in parent.paragraphs if p._p is child), None)
        elif child.tag == qn("w:tbl"):
            yield "T", next((t for t in parent.tables if t._tbl is child), None)

doc = Document(str(DOCX))
print(f"DOCX={DOCX}")
print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)}")
for idx, (kind, block) in enumerate(iter_block_items(doc), 1):
    if kind == "P":
        text = block.text.replace("\t", "\\t").replace("\n", "\\n")
        if len(text) > 180:
            text = text[:177] + "..."
        links = []
        for hyperlink in block._p.xpath('.//w:hyperlink'):
            rel_id = hyperlink.get(qn("r:id"))
            target = doc.part.rels[rel_id].target_ref if rel_id in doc.part.rels else ""
            links.append(target)
        suffix = f" links={links}" if links else ""
        print(f"{idx:03d} P style={block.style.name!r} text={text!r}{suffix}")
    else:
        rows = len(block.rows)
        cols = max((len(row.cells) for row in block.rows), default=0)
        first = " | ".join(cell.text.replace("\n", " / ") for cell in block.rows[0].cells) if rows else ""
        if len(first) > 180:
            first = first[:177] + "..."
        print(f"{idx:03d} T rows={rows} cols={cols} first={first!r}")
