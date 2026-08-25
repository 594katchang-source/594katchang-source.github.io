import html
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


sys.stdout.reconfigure(encoding="utf-8")

PROJECT = Path(__file__).resolve().parents[3]
WORK = PROJECT / "work" / "2026-08-15-seo-review-docs"
OUTPUT = WORK / "output"
SOURCE_MD = OUTPUT / "chapter-07-vitamins-seo-review.md"
REFERENCE_DOCX = OUTPUT / "chapter-05-lipids-seo-review.docx"
FINAL_DOCX = OUTPUT / "chapter-07-vitamins-seo-review.docx"

BLUE = "2E74B5"
NAVY = "1F4D78"
SLATE = "64748B"
HEADER_FILL = "E8EEF5"
TABLE_WIDTH_DXA = 9360


def set_run_font(run, size=11, color=None, bold=False, italic=False):
    run.font.name = "Calibri"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")
    rfonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def clear_document_body(doc):
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_row_rules(row, header=False):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))
    if header and tr_pr.find(qn("w:tblHeader")) is None:
        header_node = OxmlElement("w:tblHeader")
        header_node.set(qn("w:val"), "true")
        tr_pr.append(header_node)


def set_table_width(table, width_dxa=TABLE_WIDTH_DXA):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    indent = tbl_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    cell.width = Inches(width_dxa / 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def normalize_url(url):
    url = html.unescape(url.strip())
    if re.match(r"^[A-Za-z]:[\\/]", url):
        return "file:///" + url.replace("\\", "/")
    return url


def add_hyperlink(paragraph, text, url, size=11):
    hyperlink = OxmlElement("w:hyperlink")
    relationship = paragraph.part.relate_to(
        normalize_url(url),
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")
    rfonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    rpr.append(rfonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    size_node = OxmlElement("w:sz")
    size_node.set(qn("w:val"), str(int(size * 2)))
    rpr.append(size_node)
    run.append(rpr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_TOKEN = re.compile(
    r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\([^\)]+\)|<https?://[^>]+>|https?://[^\s<>]+)"
)


def add_inline(paragraph, text, size=11):
    text = html.unescape(text)
    pos = 0
    for match in INLINE_TOKEN.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=size)
        token = match.group(0)
        if token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\(([^\)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url, size=size)
        elif token.startswith("<http"):
            url = token[1:-1]
            add_hyperlink(paragraph, url, url, size=size)
        elif token.startswith("http"):
            add_hyperlink(paragraph, token.rstrip(".,)") , token.rstrip(".,)"), size=size)
            if token[-1:] in ".,)":
                run = paragraph.add_run(token[-1:])
                set_run_font(run, size=size)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size)


def add_paragraph(doc, text="", style=None, size=11):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    if text:
        add_inline(paragraph, text, size=size)
    return paragraph


def add_heading(doc, text, level):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    if level == 1:
        set_run_font(run, size=16, color=BLUE, bold=True)
    elif level == 2:
        set_run_font(run, size=13, color=BLUE, bold=True)
    else:
        set_run_font(run, size=12, color=NAVY, bold=True)
    return paragraph


def add_list_item(doc, text, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    add_inline(paragraph, text)
    return paragraph


def split_table_row(line):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [cell.strip() for cell in line.split("|")]


def is_table_separator(line):
    cells = split_table_row(line)
    return len(cells) > 1 and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def add_table(doc, rows):
    if not rows:
        return
    columns = max(len(row) for row in rows)
    table = doc.add_table(rows=0, cols=columns)
    table.style = "Table Grid"
    set_table_width(table)
    if columns == 2:
        widths = [4680, 4680]
    elif columns == 3:
        widths = [3120, 3120, 3120]
    elif columns == 4:
        widths = [2340, 2340, 2340, 2340]
    else:
        widths = [TABLE_WIDTH_DXA // columns] * columns
    for row_index, values in enumerate(rows):
        row = table.add_row()
        set_row_rules(row, header=row_index == 0)
        for col_index in range(columns):
            cell = row.cells[col_index]
            set_cell_width(cell, widths[col_index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.1
            add_inline(paragraph, values[col_index] if col_index < len(values) else "", size=9.5)
            if row_index == 0:
                set_cell_shading(cell, HEADER_FILL)
                for run in paragraph.runs:
                    run.bold = True
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)
    return table


def parse_markdown(path):
    lines = path.read_text(encoding="utf-8").splitlines()
    events = []
    title = ""
    body_started = False
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        heading = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2).strip()
            if level == 1 and not title:
                title = text
            else:
                numbered = bool(re.match(r"^\d+\.\s*", text))
                if numbered and re.match(r"^\d+\.\s*正文$", text):
                    body_started = True
                text = re.sub(r"^\d+\.\s*", "", text)
                if level == 2 and numbered:
                    document_level = 1
                elif level == 2:
                    document_level = 2
                elif level == 3 and not body_started:
                    document_level = 2
                else:
                    document_level = min(level, 3)
                events.append(("heading", document_level, text))
            i += 1
            continue
        if line.lstrip().startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            rows = [split_table_row(line)]
            i += 2
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                rows.append(split_table_row(lines[i]))
                i += 1
            events.append(("table", rows))
            continue
        list_match = re.match(r"^\s*([-*])\s+(.+)$", line)
        number_match = re.match(r"^\s*\d+[.)]\s+(.+)$", line)
        if list_match or number_match:
            numbered = bool(number_match)
            while i < len(lines):
                current = lines[i].rstrip()
                item_match = re.match(r"^\s*\d+[.)]\s+(.+)$", current) if numbered else re.match(r"^\s*[-*]\s+(.+)$", current)
                if not item_match:
                    break
                events.append(("list", item_match.group(1), numbered))
                i += 1
            continue
        block = [line.strip()]
        i += 1
        while i < len(lines):
            nxt = lines[i].rstrip()
            if not nxt.strip() or re.match(r"^#{1,6}\s+", nxt) or re.match(r"^\s*[-*]\s+", nxt) or re.match(r"^\s*\d+[.)]\s+", nxt):
                break
            if nxt.lstrip().startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
                break
            block.append(nxt.strip())
            i += 1
        events.append(("paragraph", "\n".join(block)))
    return title, events


def add_title_block(doc, title):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(title)
    set_run_font(run, size=20, color=NAVY, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Nutrition Concepts & Controversies 第 17 版書籍連載")
    set_run_font(run, size=11, color=SLATE, italic=True)
    purpose = doc.add_paragraph()
    purpose.paragraph_format.space_after = Pt(5)
    purpose.paragraph_format.line_spacing = 1.2
    label = purpose.add_run("文件用途：")
    set_run_font(label, bold=True, color=NAVY)
    value = purpose.add_run("供人工審閱的完整 SEO 草稿與研究回報，尚未代表文章通過審閱。")
    set_run_font(value)
    author = doc.add_paragraph()
    author.paragraph_format.space_after = Pt(5)
    author.paragraph_format.line_spacing = 1.2
    label = author.add_run("作者：")
    set_run_font(label, bold=True, color=NAVY)
    value = author.add_run("張雁雲營養師，Kat Chang 凱特營養師，專長為高齡營養、疾病營養、精準營養與健康促進。")
    set_run_font(value)


def build():
    if not SOURCE_MD.exists():
        raise FileNotFoundError(SOURCE_MD)
    if not REFERENCE_DOCX.exists():
        raise FileNotFoundError(REFERENCE_DOCX)
    title, events = parse_markdown(SOURCE_MD)
    doc = Document(str(REFERENCE_DOCX))
    clear_document_body(doc)
    add_title_block(doc, title)
    metadata_label = None
    in_body = False
    for event in events:
        kind = event[0]
        if kind == "heading":
            if event[2] == "正文" and event[1] == 1:
                in_body = True
                add_heading(doc, event[2], event[1])
            elif event[2] == "文章摘要" and event[1] == 2:
                metadata_label = "文章摘要"
            elif event[2] == "搜尋結果開場" and event[1] == 2:
                metadata_label = "開場"
            else:
                add_heading(doc, event[2], event[1])
        elif kind == "paragraph":
            if metadata_label:
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(5)
                paragraph.paragraph_format.line_spacing = 1.2
                label_run = paragraph.add_run(f"{metadata_label}：")
                set_run_font(label_run, bold=True, color=NAVY)
                add_inline(paragraph, event[1])
                metadata_label = None
            else:
                paragraph = add_paragraph(doc, event[1])
                if "\n" in event[1]:
                    paragraph.paragraph_format.space_after = Pt(6)
        elif kind == "list":
            if not in_body and not event[2] and event[1].startswith(("目標搜尋字詞：", "相關搜尋字詞：", "搜尋意圖：")):
                add_paragraph(doc, event[1])
            else:
                add_list_item(doc, event[1], numbered=event[2])
        elif kind == "table":
            add_table(doc, event[1])
    doc.core_properties.title = title
    doc.core_properties.subject = "Kat Chang SEO 文章審閱檔"
    doc.core_properties.author = "Kat Chang 凱特營養師"
    doc.core_properties.comments = "依前幾章 Word 審閱版型建立，供人工審閱。"
    doc.save(str(FINAL_DOCX))
    print(f"created={FINAL_DOCX}")
    print(f"title={title}")
    print(f"events={len(events)}")


if __name__ == "__main__":
    build()
