from __future__ import annotations

import os
import sys
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

ROOT = Path(__file__).resolve().parents[3]
DOCX = ROOT / "work/2026-08-15-seo-review-docs/output/chapter-08-water-minerals-seo-review.docx"
TEMP = DOCX.with_suffix(".tmp.docx")


def paragraphs(doc: Document) -> list[Paragraph]:
    return list(doc.paragraphs)


def find_one(doc: Document, predicate, label: str) -> Paragraph:
    matches = [p for p in paragraphs(doc) if predicate(p.text)]
    if len(matches) != 1:
        raise RuntimeError(f"expected one {label}, found {len(matches)}")
    return matches[0]


def replace_in_text_nodes(paragraph: Paragraph, old: str, new: str) -> int:
    hits = 0
    for node in paragraph._p.xpath(".//w:t"):
        value = node.text or ""
        if old in value:
            node.text = value.replace(old, new)
            hits += 1
    return hits


def set_plain_paragraph_text(paragraph: Paragraph, text: str) -> None:
    for child in list(paragraph._p):
        if child.tag != "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr":
            paragraph._p.remove(child)
    paragraph.add_run(text)


def main() -> None:
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)

    document = Document(DOCX)

    title_old = "每天喝多少水才夠？從電解質、鈣鐵到骨質保養的生活判讀"
    title_new = "每天喝多少水才夠？從電解質、鈣鐵到骨質保養"
    title_hits = sum(replace_in_text_nodes(p, title_old, title_new) for p in paragraphs(document))
    if title_hits != 2:
        raise RuntimeError(f"expected two title replacements, found {title_hits}")

    quick = find_one(document, lambda text: text == "省時版本：", "quick-version heading")
    following = quick._p.getnext()
    if following is None or following.tag != "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl":
        raise RuntimeError("quick-version heading is not immediately followed by its table")
    question_matches = [
        p for p in paragraphs(document)
        if p.text.startswith("本章的四個生活問題很適合拿來自我檢查：")
    ]
    if len(question_matches) != 2:
        raise RuntimeError(f"expected two body-question paragraphs, found {len(question_matches)}")
    question = question_matches[-1]
    parent = quick._p.getparent()
    parent.remove(quick._p)
    parent.remove(following)
    first_quick = (
        "省時版本：水分是日常維持循環、體溫、排便與身體功能的基礎。先看口渴、尿液、活動、天氣與疾病，"
        "再分辨一般飲水、流失過多、電解質失衡與醫師交代限水的情況。食物也能提供水分與礦物質，補充品和運動飲料不能取代對症處理。"
    )
    second_quick = (
        "從了解水分平衡與脫水開始，接著認識電解質與主要礦物質、鈣鐵等礦物質的功能，"
        "再回到三餐中的鈣來源、飲食安排、補充品與就醫警訊。"
    )
    question.insert_paragraph_before(first_quick, style="Normal")
    question.insert_paragraph_before(second_quick, style="Normal")

    water = find_one(document, lambda text: text.startswith("飲水量會隨情境改變。") and "書中把每天需要量" in text, "water context paragraph")
    water_text = (
        "飲水量會隨情境改變。高溫、濕熱、長時間戶外活動、發燒、嘔吐、腹瀉、吃較多鹽分或纖維、懷孕與哺乳，"
        "都可能改變需要。水果、蔬菜、湯、牛奶、豆漿與其他食物也含有水分，因此討論「一天喝多少」時，"
        "要先確認是在說飲品量，或是食物與飲品的總水分。"
    )
    set_plain_paragraph_text(water, water_text)
    following_water = water._p.getnext()
    if following_water is None or following_water.tag != "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p":
        raise RuntimeError("water paragraph is not followed by a paragraph")
    following_water_para = Paragraph(following_water, water._parent)
    mineral_heading = "礦物質怎麼分類？主要礦物質與微量礦物質"
    mineral_text = (
        "書中把每天需要量超過 100 mg 的鈣、氯、鎂、磷、鉀、鈉與硫酸鹽歸為主要礦物質，"
        "低於 100 mg 的碘、鐵、鋅、硒、氟、鉻、銅、錳與鉬等歸為微量礦物質。這是分類用語，微量礦物質的生理作用仍很具體。"
    )
    following_water_para.insert_paragraph_before(mineral_heading, style="Heading 2")
    following_water_para.insert_paragraph_before(mineral_text, style="Normal")

    replacements = [
        (
            "這是一個日常起點，不能直接套在心臟衰竭、腎臟病、肝硬化、使用利尿劑或醫師交代限水的人身上。",
            "這是一個日常起點建議，不適用於心臟衰竭、腎臟病、肝硬化、使用利尿劑或醫師交代限水的人身上。",
        ),
        ("口渴、尿液顏色與脫水，怎麼一起看？", "口渴、尿液顏色與脫水，怎麼看？"),
        (
            "尿液顏色可作為生活觀察線索，卻會被維生素、藥物、食物與泌尿道問題改變。只用一個訊號判斷，容易誤判。",
            "尿液顏色可作為生活觀察線索，但可能會受維生素、藥物、食物與泌尿道問題影響。只用一個訊號做判斷，容易誤判。",
        ),
        (
            "出現意識混亂、昏厥、抽搐、幾乎沒有尿、持續嘔吐、無法進水、嚴重腹瀉或疑似中暑時，這已經超出一般飲水技巧的範圍。需要迅速評估，不要只靠白開水處理。",
            "出現意識混亂、昏厥、抽搐、幾乎沒有尿、持續嘔吐、無法進水、嚴重腹瀉或疑似中暑時，問題已經超出一般飲水技巧可以處理的範圍，需要迅速接受醫療評估。這時不能只靠白開水自行補充，因為原因可能涉及嚴重脫水、電解質失衡、感染、心腎肝疾病或其他急症，補水方式與速度應由醫療團隊判斷。",
        ),
        (
            "這裡的實用原則很簡單：",
            "這裡建議的原則是：",
        ),
        ("電解質是什麼？為什麼水會跟著鹽走？", "電解質是什麼？水分分布為什麼會受溶質影響？"),
        (
            "水會受到溶液濃度與離子分布影響，這就是書中用「水會跟著鹽走」協助理解的概念。",
            "水分會在不同體液區室之間移動，受到溶液中溶質濃度與離子分布影響，因此水分與電解質要一起判讀。",
        ),
        ("單靠少撒一點鹽，常還不夠。", "不是只靠少吃食鹽、控制食鹽而已。"),
        (
            "三者功能相連，卻不代表購買一瓶「鈣鎂磷」就能補好。",
            "鈣、鎂、磷的作用彼此相關，但不能把它們當成一組買來就能一次補足。是否需要補充，仍要看飲食、檢驗、疾病與用藥，由專業人員評估。",
        ),
        ("血液中的鈣正常，不能直接證明骨骼儲備充足。", "血液中的鈣正常，不代表骨骼的鈣儲備充足。"),
        (
            "缺鐵性貧血可能帶來疲倦、頭暈、心悸、臉色變淡、活動耐受變差或注意力下降，疲倦也可能",
            "缺鐵性貧血可能帶來疲倦、頭暈、心悸、臉色變淡、活動耐受變差或注意力下降，但疲倦也可能",
        ),
        ("健康者無需把茶或乳品全部刪除", "健康者則無需避開茶或乳品"),
        ("每天吃得到的鈣來源怎麼安排？", "怎麼安排每天都吃得到鈣來源？"),
        (
            "把鈣塞進三餐，比每天睡前才想起一顆鈣片更容易持續。",
            "把鈣分布在餐食，比每天睡前才想起一顆鈣片更容易持續、益處更多。",
        ),
        ("下表是生活安排，不代表每個人的固定份量。", "下表是依據生活作息建議，不代表固定份量。"),
        ("我會把這一章轉成三個生活順序。", "水和礦物質這部分可以分成三個重點。"),
        (
            "白開水是日常基礎，分次飲用比追逐一個漂亮數字實用。",
            "補充水分是日常基礎，分次飲用比追逐某個目標數字有用。",
        ),
        ("把三天飲食、用藥、保健品、症狀與檢驗結果放在一起看，", "綜合評估三天飲食、用藥、保健品、症狀與檢驗結果，"),
        (
            "這些事情看似平常，卻比單一神奇食物更值得投入。",
            "這些事情看似很普通，卻比單一神奇食物更值得你做。",
        ),
        ("營養只能放在治療計畫裡一起安排", "營養要跟治療計畫一起安排"),
    ]
    for old, new in replacements:
        count = sum(replace_in_text_nodes(p, old, new) for p in paragraphs(document))
        if count != 1:
            raise RuntimeError(f"replacement count for {old!r}: {count}")

    document.save(TEMP)
    os.replace(TEMP, DOCX)
    print(f"updated={DOCX}")


if __name__ == "__main__":
    main()
