import sys
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding='utf-8')
ROOT = Path(__file__).resolve().parents[3]
OUTPUT = ROOT / 'work' / '2026-08-15-seo-review-docs' / 'output'
TARGETS = {
    'chapter-02-nutrition-tools-seo-review.docx': ['單一植化素', '食物基質', '我會把'],
    'chapter-06-proteins-amino-acids-seo-review.docx': ['蛋白質與胺基酸的身體', 'Controversy 6', '這張表讓我重新理解', '書中以蛋白質合成說明', '固定答案'],
    'chapter-07-vitamins-seo-review.docx': ['維生素怎麼吃才安心', '閱讀本章時', '這裡有一個很實用的分辨', '國健署第八版 DRI', '孕前與孕期葉酸是一個', '書中把補充品', '維生素的判讀順序'],
}
for filename, needles in TARGETS.items():
    doc = Document(OUTPUT / filename)
    print(f'\nFILE {filename}')
    for p in doc.paragraphs:
        if any(n in p.text for n in needles):
            print('P:', p.text[:80])
            print('RUNS:', [r.text for r in p.runs])
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if any(n in p.text for n in needles):
                        print('C:', p.text[:80])
                        print('RUNS:', [r.text for r in p.runs])
