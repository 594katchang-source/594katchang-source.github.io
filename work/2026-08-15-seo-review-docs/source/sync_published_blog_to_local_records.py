import hashlib
import json
import os
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding='utf-8')

ROOT = Path(__file__).resolve().parents[3]
POSTS_PATH = ROOT / 'blog' / 'posts.json'
WORK = ROOT / 'work' / '2026-08-15-seo-review-docs'
OUTPUT = WORK / 'output'
SOURCE = WORK / 'source'
BACKUPS = SOURCE / 'remote-sync-backups-2026-08-29'

MAPPING = {
    '2026-08-14-food-choices-human-health-guide': 'chapter-01-food-choices-seo-review.docx',
    '2026-08-15-nutrition-tools-standards-guidelines': 'chapter-02-nutrition-tools-seo-review.docx',
    '2026-08-16-remarkable-body-nutrition-guide': 'chapter-03-remarkable-body-seo-review.docx',
    '2026-08-17-carbohydrates-food-guide': 'chapter-04-carbohydrates-seo-review.docx',
    '2026-08-20-lipids-fatty-acids-guide': 'chapter-05-lipids-seo-review.docx',
    '2026-08-22-proteins-amino-acids-book-notes': 'chapter-06-proteins-amino-acids-seo-review.docx',
    '2026-08-25-vitamins-book-notes': 'chapter-07-vitamins-seo-review.docx',
}

BODY_REPLACEMENTS = {
    '2026-08-15-nutrition-tools-standards-guidelines': [
        ('單一植化素補充品有效', '單一保健補充品有效'),
        ('食物基質影響', '食物總體基質影響'),
        ('風險。\nNCCIH 提醒', '風險。\n不能忽略保健補充品也有安全風險。NCCIH 提醒'),
        ('它把營養素標準、飲食指南、食物群組、份量、食品標籤與「超級食物」放在同一條思考路線。', '同時要思考營養素標準、飲食指南、食物群組、份量、食品標籤與「超級食物」，這是同一條思考路線上。'),
        ('三個問題：\n有沒有主食來源', '三個問題：\n🔶有沒有主食來源'),
        ('目標？\n有沒有蛋白質食物與蔬菜', '目標？\n🔶有沒有蛋白質食物與蔬菜'),
        ('蔬菜？\n飲料、醬料、甜點或零食', '蔬菜？\n🔶飲料、醬料、甜點或零食'),
    ],
    '2026-08-22-proteins-amino-acids-book-notes': [
        ('蛋白質與胺基酸的身體功能、飲食品質', '蛋白質與胺基酸 從身體功能、食物品質到植物性飲食'),
        (' Controversy 6 ', ' 常見爭議 '),
        ('與 常見爭議 轉', '與常見爭議轉'),
        ('書籍重點', '重點'),
        ('我的閱讀心得', '應用'),
        ('這張表讓我重新理解', '有關'),
        ('有關「蛋白質不足」的影響。', '有關「蛋白質不足」，'),
        ('書中以蛋白質合成說明', '蛋白質合成時'),
        ('每個人的固定答案', '每個人都適合的答案'),
    ],
    '2026-08-25-vitamins-book-notes': [
        ('維生素怎麼吃才安心？從脂溶性、水溶性到補充品風險', '維生素安心吃 從脂溶性、水溶性到補充品風險'),
        ('閱讀本章時，先看', '先看'),
        ('食物是日常底盤', '食物是日常獲得營養最基本的方式'),
        ('本章從維生素的基本定義與前驅物開始', '從了解維生素的基本定義與前驅物開始'),
        ('接著整理脂溶性與水溶性', '接著認識脂溶性與水溶性的不同'),
        ('脂溶性與水溶性，先看身體如何處理', '脂溶性與水溶性，身體如何處理'),
        ('這裡有一個很實用的分辨方式', '這裡有一個很實用的知識'),
        ('把 B 群說成「補充能量」時，應該改成「協助能量代謝」', '知道B 群「不補充能量」，而是「協助能量代謝」'),
        ('才不會讓讀者以為膠囊可以替代睡眠、吃飯與休息。', '才不會讓讀者以為補充品可以替代睡眠、吃飯與休息等日常。'),
        ('「容易排出」只描述一部分生理處理', '「容易排出」只是描述一部分生理處理'),
        ('不能拿來當作高劑量安全保證', '不能拿來當作高劑量安全的保證'),
        ('別自行長期堆高', '別自行長期超量'),
        ('動物性食物可提供較容易被利用的維生素 A', '動物性食物提供的是容易被利用的維生素 A'),
        ('國健署第八版 DRI 是台灣讀者查詢參考值的主要入口，', ''),
        ('把維生素 E 直接包裝成「抗老膠囊」會漏掉劑量', '把維生素 E 當成「好用的膠囊補充品」可能會忽略劑量'),
        ('重點通常是維持穩定的飲食攝取與依醫囑監測', '應維持穩定的飲食攝取與依醫囑監測'),
        ('不能自行因為害怕凝血而把蔬菜整類刪除', '不能自行因為害怕凝血而避開某些食物'),
        ('維生素 C：膠原蛋白、鐵吸收與感冒說法', '維生素 C：膠原蛋白、鐵吸收與是否預防感冒'),
        ('蔬果是日常來源', '蔬菜水果是日常來源'),
        ('烹調水、', '烹調、'),
        ('不能只用「新鮮」兩個字判斷。', '不能只用「新鮮」兩個字判斷維生素 C高低。'),
        ('維生素 C 的問題要拆成三段：', '維生素 C 的問題要拆成三段來看：'),
        ('孕前與孕期葉酸是一個需要清楚分流的情境。', '孕前與孕期補充葉酸是女性需要注意的情形。'),
        ('這項建議要依台灣婦產科與國健署現行指引核對，不能把其他國家的補充方式直接套給每一個人。', ''),
        ('葉酸也可能遮住 B12', '葉酸也可能掩蓋住 B12'),
        ('吃肉就一定夠', '吃素需要擔心補充狀況'),
        ('書中把補充品爭議放在「需求與成本」上，這個架構很適合台灣讀者。', '補充品爭議主要是「需求與成本」上。'),
        ('補充品不能自動變成慢性病保護傘', '補充品不可能一定成為慢性病的保護傘'),
        ('也不能直接等同效果好', '也不能等同效果好'),
        ('一個可執行的補充品決策流程', '可執行的補充品決策流程'),
        ('缺口是什麼', '缺少的營養是什麼'),
        ('使用多久要回頭評估', '使用多久要重新評估'),
        ('維生素的判讀順序，是功能、來源、缺乏、過量、用藥與個人情境，單看瓶身正面很容易漏掉真正問題。', '維生素要看的是功能、來源、缺乏、過量、用藥與個人情境，單看瓶身廣告很容易漏掉真正問題。'),
        ('脂溶性與水溶性是理解吸收和儲存的入口，不能直接拿來判斷某種維生素好或壞。', '脂溶性與水溶性是用來理解吸收和儲存，不直接拿來判斷某種維生素好或壞。'),
        ('食物是日常維生素的主要底盤。', '食物是日常維生素的主要獲得方式。'),
        ('B 群不等於提神，維生素 C 不等於感冒藥，維生素 D 也不等於所有骨骼問題的單一答案。', 'B 群不一定能提神，維生素 C 不一定抗感冒，維生素 D 也不是所有骨骼問題的解決方法。'),
    ],
}

CHAPTER_07_EXTENSION = '延伸閱讀：全書導讀、第一章食物選擇、第二章營養工具與標準、第三章消化與吸收、第五章脂質與脂肪酸、作者簡介'


def sha256(path):
    return hashlib.sha256(path.read_bytes()).hexdigest()


def iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested in cell.tables:
                yield from iter_table_paragraphs(nested)


def all_paragraphs(doc):
    yield from doc.paragraphs
    for table in doc.tables:
        yield from iter_table_paragraphs(table)


def body_paragraphs(doc):
    active = False
    for child in doc.element.body.iterchildren():
        if child.tag.endswith('}p'):
            paragraph = Paragraph(child, doc)
            text = paragraph.text.strip()
            if text == '正文':
                active = True
                continue
            if active and text == 'SEO 描述':
                break
            if active:
                yield paragraph
        elif child.tag.endswith('}tbl') and active:
            table = Table(child, doc)
            for paragraph in iter_table_paragraphs(table):
                yield paragraph


def replace_in_paragraph(paragraph, replacements):
    changed = 0
    for old, new in replacements:
        if old not in paragraph.text:
            continue
        done = False
        for run in paragraph.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
                changed += 1
                done = True
        if not done:
            paragraph.text = paragraph.text.replace(old, new)
            changed += 1
    return changed


def set_next_paragraph(values, label, text):
    for index, value in enumerate(values):
        if value.strip() == label and index + 1 < len(values):
            return index + 1, text
    return None


def update_docx(post_id, post, filename):
    path = OUTPUT / filename
    before = sha256(path)
    BACKUPS.mkdir(parents=True, exist_ok=True)
    backup = BACKUPS / filename
    if not backup.exists():
        shutil.copy2(path, backup)
    doc = Document(path)
    paragraphs = list(doc.paragraphs)
    values = [paragraph.text for paragraph in paragraphs]
    changes = []

    metadata_updates = []
    title_index = next((i for i, value in enumerate(values) if value.strip() == 'SEO 標題'), None)
    if title_index is not None and title_index + 1 < len(paragraphs):
        metadata_updates.append((paragraphs[title_index + 1], post['title']))
    excerpt_index = next((i for i, value in enumerate(values) if value.strip() == '文章摘要與適合搜尋結果顯示的開場'), None)
    if excerpt_index is not None and excerpt_index + 1 < len(paragraphs):
        metadata_updates.append((paragraphs[excerpt_index + 1], '文章摘要：' + post['excerpt']))
    keyword_index = next((i for i, value in enumerate(values) if value.strip() == '目標搜尋字詞、相關搜尋字詞與搜尋意圖'), None)
    if keyword_index is not None and keyword_index + 1 < len(paragraphs):
        metadata_updates.append((paragraphs[keyword_index + 1], '目標搜尋字詞：' + '、'.join(post.get('keywords', []))))
    for paragraph in paragraphs:
        text = paragraph.text.strip()
        if text.startswith('分類：'):
            metadata_updates.append((paragraph, '分類：' + post.get('category', text[3:])))
        elif text.startswith('建議更新日期：'):
            metadata_updates.append((paragraph, '建議更新日期：' + post.get('date', text[7:])))

    for paragraph, new_text in metadata_updates:
        if paragraph.text != new_text:
            paragraph.text = new_text
            changes.append('metadata')

    replacements = BODY_REPLACEMENTS.get(post_id, [])
    for paragraph in body_paragraphs(doc):
        count = replace_in_paragraph(paragraph, replacements)
        changes.extend(['body'] * count)

    if post_id == '2026-08-25-vitamins-book-notes':
        current_body_texts = [p.text.strip() for p in body_paragraphs(doc)]
        if CHAPTER_07_EXTENSION not in current_body_texts:
            seo_heading = next((p for p in doc.paragraphs if p.text.strip() == 'SEO 描述'), None)
            if seo_heading is not None:
                new_paragraph = doc.add_paragraph(CHAPTER_07_EXTENSION)
                seo_heading._p.addprevious(new_paragraph._p)
                changes.append('body-extension')

    tmp = path.with_suffix('.remote-sync.tmp.docx')
    doc.save(tmp)
    os.replace(tmp, path)
    after = sha256(path)
    print(f'DOCX {filename} changes={len(changes)} before={before} after={after}')


def update_source_snapshot(post_id, post, json_name, html_name):
    json_path = SOURCE / json_name
    html_path = SOURCE / html_name
    for path in (json_path, html_path):
        backup = BACKUPS / path.name
        if path.exists() and not backup.exists():
            shutil.copy2(path, backup)
    snapshot = json.loads(json_path.read_text(encoding='utf-8')) if json_path.exists() else {}
    for key, value in post.items():
        snapshot[key] = value
    json_path.write_text(json.dumps(snapshot, ensure_ascii=False, indent=2) + '\n', encoding='utf-8')
    html_path.write_text(post['body'].rstrip() + '\n', encoding='utf-8')
    print(f'SOURCE {json_name} and {html_name} synced to {post_id}')


posts = json.loads(POSTS_PATH.read_text(encoding='utf-8'))['posts']
post_by_id = {post['id']: post for post in posts}
for post_id, filename in MAPPING.items():
    update_docx(post_id, post_by_id[post_id], filename)

update_source_snapshot(
    '2026-08-22-proteins-amino-acids-book-notes',
    post_by_id['2026-08-22-proteins-amino-acids-book-notes'],
    'chapter-06-publish.json',
    'chapter-06-publish.html',
)
update_source_snapshot(
    '2026-08-25-vitamins-book-notes',
    post_by_id['2026-08-25-vitamins-book-notes'],
    'chapter-07-publish.json',
    'chapter-07-publish.html',
)
