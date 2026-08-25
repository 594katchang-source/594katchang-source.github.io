import json
import re
from pathlib import Path
from zipfile import ZipFile
from xml.etree import ElementTree as ET

from docx import Document


BASE = Path(__file__).resolve().parent.parent
SOURCE = BASE / "source" / "chapter-04-review.json"
HTML = BASE / "source" / "chapter-04-review.html"
DOCX = BASE / "output" / "chapter-04-carbohydrates-seo-review.docx"
NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def main():
    data = json.loads(SOURCE.read_text(encoding="utf-8"))
    html = HTML.read_text(encoding="utf-8")
    assert DOCX.exists() and DOCX.stat().st_size > 0
    assert data["seoTitle"] == "碳水化合物怎麼吃才穩？從全穀、膳食纖維到添加糖"
    assert "本篇整理書籍" in data["bodyHtml"] and "Nutrition Concepts &amp; Controversies" in data["bodyHtml"]
    assert "本篇章節" in data["bodyHtml"] and "The Carbohydrates: Sugar, Starch, Glycogen, and Fiber" in data["bodyHtml"]
    assert "文章性質" in data["bodyHtml"] and "章節整理發文" in data["bodyHtml"]
    assert any(
        marker in data["bodyHtml"]
        for marker in [
            "<h2>糖尿病安全段落：先確認藥物與血糖監測，再調整醣類</h2>",
            "<h2>糖尿病安全作法：先確認藥物與血糖監測，再調整醣類</h2>",
        ]
    )
    assert "<h2>七日練習表：把醣類判讀變成日常習慣</h2>" in data["bodyHtml"]
    assert "<h2>Kat Chang 營養師的判讀</h2>" in data["bodyHtml"]
    assert data["wordCount"]["characters"] >= 2000
    assert data["bodyHtml"] in html
    assert f'<link rel="canonical" href="{data["canonical"]}">' in html
    assert len(re.findall(r"<h2>", data["bodyHtml"])) == 13
    assert len(re.findall(r"<h3>", data["bodyHtml"])) == 5
    assert len(re.findall(r"<table>", data["bodyHtml"])) == 10
    for label, url in data["internalLinks"]:
        assert url in data["bodyHtml"] and url in html, url
    for item in data["faqEntities"]:
        assert item["question"] in data["bodyHtml"] and item["answer"] in data["bodyHtml"]

    with ZipFile(DOCX) as zf:
        assert "word/document.xml" in zf.namelist()
        root = ET.fromstring(zf.read("word/document.xml"))
        rel_root = ET.fromstring(zf.read("word/_rels/document.xml.rels"))
    rel_map = {item.get("Id"): item.get("Target") for item in rel_root}
    hyperlink_targets = []
    for hyperlink in root.findall(".//w:hyperlink", NS):
        rid = hyperlink.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
        if rid and rel_map.get(rid):
            hyperlink_targets.append(rel_map[rid])
    assert len(hyperlink_targets) == len(data["internalLinks"])
    for _, url in data["internalLinks"]:
        assert url in hyperlink_targets
    document = Document(DOCX)
    paragraph_text = "\n".join(p.text for p in document.paragraphs)
    assert data["seoTitle"] in paragraph_text
    assert "Kat Chang 營養師的判讀" in paragraph_text
    assert "七日練習表：把醣類判讀變成日常習慣" in paragraph_text
    assert any(
        marker in paragraph_text
        for marker in [
            "糖尿病安全段落：先確認藥物與血糖監測，再調整醣類",
            "糖尿病安全作法：先確認藥物與血糖監測，再調整醣類",
        ]
    )
    assert len(document.tables) >= 11

    table_audit = []
    for table in root.findall(".//w:tbl", NS):
        tbl_pr = table.find("./w:tblPr", NS)
        width = tbl_pr.find("./w:tblW", NS)
        layout = tbl_pr.find("./w:tblLayout", NS)
        assert width is not None and width.get(f"{{{NS['w']}}}w") == "9360"
        assert layout is not None and layout.get(f"{{{NS['w']}}}type") == "fixed"
        rows = table.findall("./w:tr", NS)
        assert rows and rows[0].find("./w:trPr/w:tblHeader", NS) is not None
        for row in rows:
            assert row.find("./w:trPr/w:cantSplit", NS) is not None
        table_audit.append(len(rows))

    print(json.dumps({
        "source": str(SOURCE),
        "html": str(HTML),
        "docx": str(DOCX),
        "bodyCharacters": data["wordCount"]["characters"],
        "bodyTables": len(re.findall(r"<table>", data["bodyHtml"])),
        "bodyH2": len(re.findall(r"<h2>", data["bodyHtml"])),
        "bodyH3": len(re.findall(r"<h3>", data["bodyHtml"])),
        "faq": len(data["faqEntities"]),
        "internalLinks": len(data["internalLinks"]),
        "docxParagraphs": len(document.paragraphs),
        "docxTables": len(document.tables),
        "docxTableRows": table_audit,
        "sameSourceBody": True,
        "zipValid": True,
        "presetGeometry": "9360 DXA, fixed layout, header rows, cantSplit",
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
