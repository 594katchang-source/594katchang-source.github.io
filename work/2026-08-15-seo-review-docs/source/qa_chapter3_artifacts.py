import json
import re
import zipfile
from html import unescape
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


BASE = Path(__file__).resolve().parent.parent
SOURCE_DIR = BASE / "source"
OUTPUT_DIR = BASE / "output"
PROJECT_ROOT = BASE


FORBIDDEN = [
    "再者", "然而", "不過", "值得注意的是", "由此可見", "可以看出", "總之", "總而言之",
    "總的來說", "整體而言", "大體而言", "綜上所述", "綜合來看", "綜觀以上", "換句話說",
    "換言之", "另一方面", "相對而言", "與此同時", "首先", "其次", "再次", "最後",
    "不可否認的是", "毋庸置疑的是", "全面", "全方位", "系統性", "深度", "多層次",
    "多維度", "多角度", "顛覆性", "革命性", "劃時代", "前所未有", "意義重大",
    "具有重要意義", "影響深遠", "不可忽視", "舉足輕重", "令人印象深刻", "發人深省",
    "耐人尋味", "引人入勝", "耳目一新", "極具潛力", "極具前景", "極具發展空間",
    "不僅", "不只", "不是", "而是", "接下來", "我們將探討", "讓我們一起來看看",
    "這篇文章將帶你了解", "透過上述介紹", "相信你已經", "無論你是", "如果你也有同樣的困惑"
]


def visible_text(markup):
    text = re.sub(r"<script[\s\S]*?</script>", " ", markup, flags=re.I)
    text = re.sub(r"<style[\s\S]*?</style>", " ", text, flags=re.I)
    text = re.sub(r"<[^>]+>", " ", text)
    return re.sub(r"\s+", " ", unescape(text)).strip()


def table_widths(table):
    widths = []
    for cell in table.rows[0].cells:
        tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
        widths.append(int(tc_w.get(qn("w:w"))) if tc_w is not None else 0)
    return widths


def main():
    data_path = SOURCE_DIR / "chapter-03-review.json"
    html_path = SOURCE_DIR / "chapter-03-review.html"
    docx_path = OUTPUT_DIR / "chapter-03-remarkable-body-seo-review.docx"
    source_path = Path(r"D:\@Codex\書籍\2026-07-29-Nutrition-Concepts-Controversies-17e\process\chapter-03-source.txt")

    data = json.loads(data_path.read_text(encoding="utf-8"))
    html_text = html_path.read_text(encoding="utf-8")
    body = visible_text(data["bodyHtml"])
    assert data["bodyHtml"] in html_text, "HTML body differs from JSON body source"
    assert data["wordCount"]["characters"] == len(body), "wordCount characters mismatch"
    assert len(body) >= 2000, f"body too short: {len(body)}"
    assert len(data["faq"]) == 5, "FAQ count must remain five"
    assert len(data["internalLinks"]) >= 5, "internal link suggestions are incomplete"
    assert data["articleSchema"].count("showOnHome: false") == 1, "new article must default showOnHome false"

    review_copy = "\n".join([
        data["seoTitle"], data["searchIntent"], data["summary"], data["opening"],
        data["articleTitle"], data["bodyHtml"], data["seoDescription"], data["category"],
        *data["tags"], *data["faq"], *data["originalClaims"], *data["pending"]
    ])
    review_copy = re.sub(r"https?://[^\s\"'<>]+", " ", review_copy)
    found_forbidden = [term for term in FORBIDDEN if term in review_copy]
    assert not found_forbidden, f"forbidden wording found: {found_forbidden}"

    source_text = source_path.read_text(encoding="utf-8")
    assert "Chapter 3 The Remarkable Body" in source_text
    assert source_text.count("===== PDF PAGE") >= 30
    assert "�" not in source_text
    assert "???" not in source_text

    with zipfile.ZipFile(docx_path) as archive:
        assert archive.testzip() is None, "DOCX ZIP test failed"
    doc = Document(docx_path)
    assert len(doc.paragraphs) >= 100, f"paragraph count too low: {len(doc.paragraphs)}"
    assert len(doc.tables) >= 10, f"table count too low: {len(doc.tables)}"
    table_failures = []
    cant_split_failures = []
    header_failures = []
    for index, table in enumerate(doc.tables, start=1):
        if sum(table_widths(table)) != 9360:
            table_failures.append((index, table_widths(table)))
        for row_index, row in enumerate(table.rows):
            tr_pr = row._tr.get_or_add_trPr()
            if tr_pr.find(qn("w:cantSplit")) is None:
                cant_split_failures.append((index, row_index))
            if row_index == 0 and tr_pr.find(qn("w:tblHeader")) is None:
                header_failures.append(index)
    assert not table_failures, table_failures
    assert not cant_split_failures, cant_split_failures
    assert not header_failures, header_failures

    print(json.dumps({
        "bodyCharacters": len(body),
        "bodyWords": len(body.split()),
        "h2": len(re.findall(r"<h2>", data["bodyHtml"])),
        "h3": len(re.findall(r"<h3>", data["bodyHtml"])),
        "bodyTables": len(re.findall(r"<table>", data["bodyHtml"])),
        "faq": len(data["faq"]),
        "internalLinks": len(data["internalLinks"]),
        "sourcePageMarkers": source_text.count("===== PDF PAGE"),
        "docxParagraphs": len(doc.paragraphs),
        "docxTables": len(doc.tables),
        "docxZip": "pass",
        "sameSource": "pass",
        "forbiddenWording": "pass",
        "chapter3Published": False
    }, ensure_ascii=False))


if __name__ == "__main__":
    main()
