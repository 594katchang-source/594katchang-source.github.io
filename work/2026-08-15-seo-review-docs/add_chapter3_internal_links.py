from pathlib import Path

from docx import Document

from build_review_docs import add_hyperlink


BASE = Path(__file__).resolve().parent
DOCX_PATH = BASE / "output" / "chapter-03-remarkable-body-seo-review.docx"

LINKS = [
    ("全書導讀", "https://594katchang-source.github.io/blog/post.html?id=2026-08-13-nutrition-concepts-controversies-17e-guide"),
    ("第二章：Nutrition Tools, Standards, and Guidelines", "https://594katchang-source.github.io/blog/post.html?id=2026-08-15-nutrition-tools-standards-guidelines"),
    ("第一章：Food Choices and Human Health", "https://594katchang-source.github.io/blog/post.html?id=2026-08-14-food-choices-human-health-guide"),
    ("功能醫學與預防阿茲海默症的飲食介入判讀", "https://594katchang-source.github.io/blog/post.html?id=2026-05-19-功能醫學預防阿茲海默症的系統性介入策略"),
    ("網站首頁", "https://594katchang-source.github.io/"),
]


def main():
    doc = Document(DOCX_PATH)
    marker = "FAQ：人體如何處理食物與營養？"
    if any(p.text.strip().startswith("延伸閱讀：") for p in doc.paragraphs):
        print("already-present")
        return
    target = next((p for p in doc.paragraphs if p.text.strip() == marker), None)
    if target is None:
        raise RuntimeError("FAQ heading not found")
    paragraph = target.insert_paragraph_before()
    paragraph.add_run("延伸閱讀：")
    for index, (label, url) in enumerate(LINKS):
        if index:
            paragraph.add_run("、")
        add_hyperlink(paragraph, label, url)
    doc.save(DOCX_PATH)
    print({"added": len(LINKS), "paragraph": paragraph.text})


if __name__ == "__main__":
    main()
